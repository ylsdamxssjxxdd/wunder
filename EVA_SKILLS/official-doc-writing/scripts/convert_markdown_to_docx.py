#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import argparse
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DEFAULT_BODY_FONT = "仿宋GB2312"
DEFAULT_TITLE_FONT = "方正小标宋简体"
DEFAULT_HEADING1_FONT = "黑体"
DEFAULT_HEADING2_FONT = "楷体GB2312"
DEFAULT_HEADING3_FONT = "仿宋GB2312"
DEFAULT_HEADING4_FONT = "仿宋GB2312"
DEFAULT_DIGIT_FONT = "Times New Roman"
DEFAULT_PAGE_NUMBER_FONT = "宋体"
DEFAULT_CODE_FONT = "Courier New"


class PandocNotFoundError(RuntimeError):
    pass


class PandocFailedError(RuntimeError):
    pass


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert Markdown to DOCX using official document layout rules.",
    )
    parser.add_argument("input", help="Path to Markdown file.")
    parser.add_argument(
        "--output",
        default="",
        help="Output DOCX path (default: same name as input).",
    )
    parser.add_argument("--font", default=DEFAULT_BODY_FONT, help="Body font name.")
    parser.add_argument("--font-size", type=float, default=16, help="Body font size (pt).")
    parser.add_argument("--title-font", default=DEFAULT_TITLE_FONT, help="Title font name.")
    parser.add_argument("--title-size", type=float, default=22, help="Title font size (pt).")
    parser.add_argument(
        "--heading1-font",
        default=DEFAULT_HEADING1_FONT,
        help="Level-1 heading font.",
    )
    parser.add_argument(
        "--heading2-font",
        default=DEFAULT_HEADING2_FONT,
        help="Level-2 heading font.",
    )
    parser.add_argument(
        "--heading3-font",
        default=DEFAULT_HEADING3_FONT,
        help="Level-3 heading font.",
    )
    parser.add_argument(
        "--heading4-font",
        default=DEFAULT_HEADING4_FONT,
        help="Level-4 heading font.",
    )
    parser.add_argument("--heading-size", type=float, default=16, help="Heading font size.")
    parser.add_argument("--digit-font", default=DEFAULT_DIGIT_FONT, help="Digit font name.")
    parser.add_argument("--code-font", default=DEFAULT_CODE_FONT, help="Inline code font name.")
    parser.add_argument(
        "--page-number-font",
        default=DEFAULT_PAGE_NUMBER_FONT,
        help="Page number font name.",
    )
    parser.add_argument(
        "--page-number-size",
        type=float,
        default=14,
        help="Page number font size (pt).",
    )
    parser.add_argument(
        "--line-spacing-pt",
        type=float,
        default=28.9,
        help="Fixed line spacing in points.",
    )
    parser.add_argument(
        "--chars-per-line",
        type=int,
        default=28,
        help="Characters per line used to build doc grid.",
    )
    parser.add_argument(
        "--no-doc-grid",
        action="store_true",
        help="Disable document grid setup.",
    )
    parser.add_argument(
        "--no-page-number",
        action="store_true",
        help="Disable page number insertion.",
    )
    parser.add_argument("--margin-top-cm", type=float, default=3.7, help="Top margin in cm.")
    parser.add_argument(
        "--margin-bottom-cm",
        type=float,
        default=3.5,
        help="Bottom margin in cm.",
    )
    parser.add_argument("--margin-left-cm", type=float, default=2.8, help="Left margin in cm.")
    parser.add_argument(
        "--margin-right-cm",
        type=float,
        default=2.6,
        help="Right margin in cm.",
    )
    parser.add_argument(
        "--first-line-indent-cm",
        type=float,
        default=0.0,
        help="First line indent in cm (default: 0).",
    )
    parser.add_argument(
        "--first-line-indent-chars",
        type=float,
        default=2.0,
        help="First line indent in chars when cm is 0 (default: 2).",
    )
    parser.add_argument(
        "--reference-doc",
        default="",
        help="Optional reference DOCX template. If provided, skip auto template.",
    )
    parser.add_argument(
        "--use-pandoc",
        action="store_true",
        help="Use pandoc for conversion (default: internal converter).",
    )
    return parser.parse_args()


def ensure_rfonts(element, east_asia_font: str, ascii_font: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(
    style,
    font_name: str,
    font_size: Optional[float],
    bold: Optional[bool],
    digit_font: str,
) -> None:
    style.font.name = font_name
    ensure_rfonts(style.element, font_name, digit_font)
    if font_size is not None:
        style.font.size = Pt(font_size)
    if bold is not None:
        style.font.bold = bold


def set_run_font(
    run,
    font_name: str,
    font_size: Optional[float],
    bold: Optional[bool],
    digit_font: str,
) -> None:
    run.font.name = font_name
    ensure_rfonts(run._element, font_name, digit_font)
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold


def configure_section(section, args: argparse.Namespace) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(args.margin_top_cm)
    section.bottom_margin = Cm(args.margin_bottom_cm)
    section.left_margin = Cm(args.margin_left_cm)
    section.right_margin = Cm(args.margin_right_cm)

    if args.no_doc_grid:
        return

    if args.chars_per_line <= 0:
        return

    text_width_cm = 21.0 - args.margin_left_cm - args.margin_right_cm
    text_width_pt = text_width_cm / 2.54 * 72
    char_space = int(text_width_pt / args.chars_per_line * 20)
    line_pitch = int(args.line_spacing_pt * 20)

    sect_pr = section._sectPr
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    doc_grid.set(qn("w:type"), "linesAndChars")
    doc_grid.set(qn("w:linePitch"), str(line_pitch))
    doc_grid.set(qn("w:charSpace"), str(char_space))


def setup_styles(doc: Document, args: argparse.Namespace) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, args.font, args.font_size, None, args.digit_font)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    if args.first_line_indent_cm > 0:
        normal.paragraph_format.first_line_indent = Cm(args.first_line_indent_cm)
    else:
        normal.paragraph_format.first_line_indent = Pt(
            args.font_size * args.first_line_indent_chars
        )

    heading_1 = doc.styles["Heading 1"]
    set_style_font(heading_1, args.title_font, args.title_size, False, args.digit_font)
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(0)
    heading_1.paragraph_format.first_line_indent = Pt(0)
    heading_1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_1.paragraph_format.line_spacing = Pt(args.line_spacing_pt)

    heading_2 = doc.styles["Heading 2"]
    set_style_font(heading_2, args.heading1_font, args.heading_size, False, args.digit_font)
    heading_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_2.paragraph_format.space_before = Pt(0)
    heading_2.paragraph_format.space_after = Pt(0)
    heading_2.paragraph_format.first_line_indent = Pt(0)
    heading_2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_2.paragraph_format.line_spacing = Pt(args.line_spacing_pt)

    heading_3 = doc.styles["Heading 3"]
    set_style_font(heading_3, args.heading2_font, args.heading_size, False, args.digit_font)
    heading_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_3.paragraph_format.space_before = Pt(0)
    heading_3.paragraph_format.space_after = Pt(0)
    heading_3.paragraph_format.first_line_indent = Pt(0)
    heading_3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_3.paragraph_format.line_spacing = Pt(args.line_spacing_pt)

    heading_4 = doc.styles["Heading 4"]
    set_style_font(heading_4, args.heading3_font, args.heading_size, False, args.digit_font)
    heading_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_4.paragraph_format.space_before = Pt(0)
    heading_4.paragraph_format.space_after = Pt(0)
    heading_4.paragraph_format.first_line_indent = Pt(0)
    heading_4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_4.paragraph_format.line_spacing = Pt(args.line_spacing_pt)

    heading_5 = doc.styles["Heading 5"]
    set_style_font(heading_5, args.heading4_font, args.heading_size, False, args.digit_font)
    heading_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_5.paragraph_format.space_before = Pt(0)
    heading_5.paragraph_format.space_after = Pt(0)
    heading_5.paragraph_format.first_line_indent = Pt(0)
    heading_5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_5.paragraph_format.line_spacing = Pt(args.line_spacing_pt)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        set_style_font(style, args.font, args.font_size, None, args.digit_font)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    if "Imprint" not in doc.styles:
        imprint = doc.styles.add_style("Imprint", WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(imprint, args.font, 14, None, args.digit_font)
        imprint.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        imprint.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
        imprint.paragraph_format.space_before = Pt(0)
        imprint.paragraph_format.space_after = Pt(0)
        imprint.paragraph_format.first_line_indent = Pt(0)


def add_page_field(paragraph, font_name: str, font_size: float) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_pr.append(r_fonts)

    size_value = str(int(font_size * 2))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size_value)
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), size_value)
    r_pr.append(sz_cs)

    run.append(r_pr)
    field.append(run)
    paragraph._p.append(field)


def add_page_numbers(doc: Document, section, args: argparse.Namespace) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False

    def ensure_paragraph(footer, alignment):
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        add_page_field(paragraph, args.page_number_font, args.page_number_size)

    ensure_paragraph(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
    ensure_paragraph(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)


def create_base_document(args: argparse.Namespace) -> Document:
    doc = Document()
    section = doc.sections[0]
    configure_section(section, args)
    setup_styles(doc, args)
    if not args.no_page_number:
        add_page_numbers(doc, section, args)
    return doc


def build_reference_docx(path: Path, args: argparse.Namespace) -> None:
    doc = create_base_document(args)
    doc.save(path)


CHINESE_DIGITS = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]


def to_chinese_number(value: int) -> str:
    if value <= 0:
        return str(value)
    if value < 10:
        return CHINESE_DIGITS[value]
    if value == 10:
        return "十"
    if value < 20:
        return f"十{CHINESE_DIGITS[value % 10]}"
    if value < 100:
        tens = CHINESE_DIGITS[value // 10]
        if value % 10 == 0:
            return f"{tens}十"
        return f"{tens}十{CHINESE_DIGITS[value % 10]}"
    return str(value)


LEVEL2_NUMBER_RE = re.compile(r"^([一二三四五六七八九十百千]+、|\d+[\.、])")
LEVEL3_NUMBER_RE = re.compile(r"^[（(][一二三四五六七八九十百千]+[）)]")
LEVEL4_NUMBER_RE = re.compile(r"^\d+[\.、]")
LEVEL5_NUMBER_RE = re.compile(r"^[（(]\d+[）)]")
TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-+:?$")
LINE_BREAK_RE = re.compile(r"(\\\\)\\s*$", re.IGNORECASE)
BR_TAG_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
ASCII_WORD_RE = re.compile(r"[A-Za-z0-9]")


def has_numbering(level: int, title: str) -> bool:
    if level == 2:
        return bool(LEVEL2_NUMBER_RE.match(title))
    if level == 3:
        return bool(LEVEL3_NUMBER_RE.match(title))
    if level == 4:
        return bool(LEVEL4_NUMBER_RE.match(title))
    if level == 5:
        return bool(LEVEL5_NUMBER_RE.match(title))
    return False


def apply_heading_numbering(md_text: str) -> str:
    lines = md_text.splitlines()
    counters = [0, 0, 0, 0]
    output = []
    in_code_block = False
    heading_re = re.compile(r"^(#{1,5})\s+(.*)$")
    title_count = 0

    for raw in lines:
        line = raw.rstrip("\n")
        stripped = line.lstrip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            output.append(line)
            continue
        if in_code_block:
            output.append(line)
            continue

        match = heading_re.match(stripped)
        if not match:
            output.append(line)
            continue

        prefix, title = match.groups()
        level = len(prefix)
        if not title.strip():
            output.append(line)
            continue

        if level == 1:
            title_count += 1
            counters = [0, 0, 0, 0]
            numbered = title
        elif level == 2:
            counters[0] += 1
            counters[1:] = [0, 0, 0]
            numbered = (
                title
                if has_numbering(level, title)
                else f"{to_chinese_number(counters[0])}、{title}"
            )
        elif level == 3:
            counters[1] += 1
            counters[2:] = [0, 0]
            numbered = (
                title
                if has_numbering(level, title)
                else f"（{to_chinese_number(counters[1])}）{title}"
            )
        elif level == 4:
            counters[2] += 1
            counters[3] = 0
            numbered = (
                title if has_numbering(level, title) else f"{counters[2]}.{title}"
            )
        else:
            counters[3] += 1
            numbered = (
                title if has_numbering(level, title) else f"（{counters[3]}）{title}"
            )

        indent = line[: len(line) - len(stripped)]
        output.append(f"{indent}{prefix} {numbered}")

    if title_count > 1:
        raise ValueError("Only one top-level '#' title is allowed.")

    return "\n".join(output)


def add_red_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF0000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_right_paragraph(doc: Document, text: str, args: argparse.Namespace) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    add_text_with_breaks(paragraph, text, args, bold=None)


def set_table_cell_margins(
    table,
    top_cm: float,
    bottom_cm: float,
    left_cm: float,
    right_cm: float,
) -> None:
    def to_twips(value_cm: float) -> int:
        return int(value_cm / 2.54 * 1440)

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for tag, value in (
        ("top", top_cm),
        ("bottom", bottom_cm),
        ("left", left_cm),
        ("right", right_cm),
    ):
        node = tbl_cell_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(to_twips(value)))
        node.set(qn("w:type"), "dxa")


def apply_run_format(
    run,
    args: argparse.Namespace,
    bold: Optional[bool],
    italic: Optional[bool],
    strike: Optional[bool],
    font_name: Optional[str],
    font_size: Optional[float],
    code: bool = False,
) -> None:
    if code:
        set_run_font(
            run,
            args.code_font,
            font_size if font_size is not None else args.font_size,
            False,
            args.code_font,
        )
        run.font.italic = False
        run.font.strike = False
        return

    if font_name is not None or font_size is not None or bold is not None:
        set_run_font(
            run,
            font_name if font_name is not None else args.font,
            font_size if font_size is not None else args.font_size,
            bold,
            args.digit_font,
        )
    elif bold is not None:
        run.font.bold = bold

    if italic:
        run.font.italic = True
    if strike:
        run.font.strike = True


def add_inline_runs(
    paragraph,
    text: str,
    args: argparse.Namespace,
    bold: Optional[bool],
    font_name: Optional[str],
    font_size: Optional[float],
) -> bool:
    added = False
    for token in parse_inline_markdown(text):
        if not token.text:
            continue
        run = paragraph.add_run(token.text)
        if token.code:
            apply_run_format(
                run,
                args,
                bold=False,
                italic=False,
                strike=False,
                font_name=args.code_font,
                font_size=font_size,
                code=True,
            )
        else:
            apply_run_format(
                run,
                args,
                bold=bold,
                italic=token.italic if token.italic else None,
                strike=token.strike if token.strike else None,
                font_name=font_name,
                font_size=font_size,
            )
            if token.bold:
                run.font.bold = True
        added = True
    return added


def add_text_with_breaks(
    paragraph,
    text: str,
    args: argparse.Namespace,
    bold: Optional[bool],
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
) -> None:
    for segment, inline_break in split_inline_breaks(text):
        segment_text, hard_break = extract_manual_break(segment)
        segment_text = segment_text.strip()
        if segment_text:
            add_inline_runs(
                paragraph,
                segment_text,
                args,
                bold=bold,
                font_name=font_name,
                font_size=font_size,
            )
        if hard_break or inline_break:
            run = paragraph.add_run("")
            apply_run_format(
                run,
                args,
                bold=bold,
                italic=None,
                strike=None,
                font_name=font_name,
                font_size=font_size,
            )
            run.add_break()


def add_code_block(doc: Document, lines: list[str], args: argparse.Namespace) -> None:
    if not lines:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
    for index, line in enumerate(lines):
        if index > 0:
            run = paragraph.add_run("")
            apply_run_format(
                run,
                args,
                bold=None,
                italic=None,
                strike=None,
                font_name=args.code_font,
                font_size=args.font_size,
                code=True,
            )
            run.add_break()
        run = paragraph.add_run(line)
        apply_run_format(
            run,
            args,
            bold=None,
            italic=None,
            strike=None,
            font_name=args.code_font,
            font_size=args.font_size,
            code=True,
        )


def needs_space(prev: str, next_text: str) -> bool:
    if not prev or not next_text:
        return False
    prev_char = prev[-1]
    next_char = next_text[0]
    if ASCII_WORD_RE.match(prev_char) and ASCII_WORD_RE.match(next_char):
        return True
    if prev_char in ",.;:!?" and next_char.isalnum():
        return True
    return False


def extract_manual_break(raw_line: str) -> tuple[str, bool]:
    if raw_line.endswith("  "):
        return raw_line.rstrip(), True
    stripped = raw_line.rstrip()
    match = LINE_BREAK_RE.search(stripped)
    if match:
        return stripped[: match.start()].rstrip(), True
    return stripped, False


def split_table_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def is_table_row(line: str) -> bool:
    if "|" not in line:
        return False
    cells = split_table_row(line)
    return len(cells) >= 2


def is_table_separator_line(line: str) -> bool:
    if "|" not in line:
        return False
    cells = split_table_row(line)
    if len(cells) < 2:
        return False
    return all(TABLE_SEPARATOR_CELL_RE.match(cell.strip()) for cell in cells)


def parse_table_alignments(line: str) -> list[WD_ALIGN_PARAGRAPH]:
    alignments: list[WD_ALIGN_PARAGRAPH] = []
    cells = split_table_row(line)
    for cell in cells:
        cell_text = cell.strip()
        if not TABLE_SEPARATOR_CELL_RE.match(cell_text):
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            continue
        left = cell_text.startswith(":")
        right = cell_text.endswith(":")
        if left and right:
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif right:
            alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
    return alignments


def split_inline_breaks(raw_line: str) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    last_index = 0
    for match in BR_TAG_RE.finditer(raw_line):
        segment = raw_line[last_index:match.start()]
        segments.append((segment, True))
        last_index = match.end()
    segments.append((raw_line[last_index:], False))
    return segments


@dataclass
class InlineStyle:
    bold: bool = False
    italic: bool = False
    strike: bool = False

    def copy(self) -> "InlineStyle":
        return InlineStyle(self.bold, self.italic, self.strike)


@dataclass
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False


def parse_inline_markdown(text: str, base_style: Optional[InlineStyle] = None) -> list[InlineRun]:
    runs: list[InlineRun] = []
    buffer: list[str] = []
    state = base_style.copy() if base_style else InlineStyle()

    def flush_buffer() -> None:
        if not buffer:
            return
        runs.append(
            InlineRun(
                "".join(buffer),
                bold=state.bold,
                italic=state.italic,
                strike=state.strike,
            )
        )
        buffer.clear()

    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "\\":
            if i + 1 < len(text):
                buffer.append(text[i + 1])
                i += 2
                continue
            buffer.append(ch)
            i += 1
            continue
        if ch == "`":
            tick_count = 1
            while i + tick_count < len(text) and text[i + tick_count] == "`":
                tick_count += 1
            marker = "`" * tick_count
            end_index = text.find(marker, i + tick_count)
            if end_index != -1:
                flush_buffer()
                code_text = text[i + tick_count : end_index]
                if code_text:
                    runs.append(InlineRun(code_text, code=True))
                i = end_index + tick_count
                continue
        if ch == "!" and i + 1 < len(text) and text[i + 1] == "[":
            image = parse_markdown_link(text, i + 1)
            if image is not None:
                alt_text, url, end_index = image
                flush_buffer()
                runs.append(
                    InlineRun(
                        "图片：",
                        bold=state.bold,
                        italic=state.italic,
                        strike=state.strike,
                    )
                )
                if alt_text:
                    runs.extend(parse_inline_markdown(alt_text, state.copy()))
                if url:
                    runs.append(
                        InlineRun(
                            f"（{url}）",
                            bold=state.bold,
                            italic=state.italic,
                            strike=state.strike,
                        )
                    )
                i = end_index
                continue
        if ch == "[":
            link = parse_markdown_link(text, i)
            if link is not None:
                link_text, url, end_index = link
                flush_buffer()
                if link_text:
                    runs.extend(parse_inline_markdown(link_text, state.copy()))
                if url:
                    runs.append(
                        InlineRun(
                            f"（{url}）",
                            bold=state.bold,
                            italic=state.italic,
                            strike=state.strike,
                        )
                    )
                i = end_index
                continue
        if ch == "<":
            autolink = parse_autolink(text, i)
            if autolink is not None:
                url, end_index = autolink
                flush_buffer()
                runs.append(
                    InlineRun(
                        url,
                        bold=state.bold,
                        italic=state.italic,
                        strike=state.strike,
                    )
                )
                i = end_index
                continue

        marker_handled = False
        for marker, toggles in (
            ("***", ("bold", "italic")),
            ("___", ("bold", "italic")),
            ("**", ("bold",)),
            ("__", ("bold",)),
            ("*", ("italic",)),
            ("_", ("italic",)),
            ("~~", ("strike",)),
        ):
            if not text.startswith(marker, i):
                continue
            if not can_toggle_marker(text, i, marker):
                break
            if text.find(marker, i + len(marker)) == -1:
                break
            flush_buffer()
            for flag in toggles:
                if flag == "bold":
                    state.bold = not state.bold
                elif flag == "italic":
                    state.italic = not state.italic
                elif flag == "strike":
                    state.strike = not state.strike
            i += len(marker)
            marker_handled = True
            break
        if marker_handled:
            continue

        buffer.append(ch)
        i += 1

    flush_buffer()
    return runs


def parse_markdown_link(text: str, start: int) -> Optional[tuple[str, str, int]]:
    if start >= len(text) or text[start] != "[":
        return None
    end_label = find_matching_bracket(text, start, "[", "]")
    if end_label is None:
        return None
    if end_label + 1 >= len(text) or text[end_label + 1] != "(":
        return None
    end_url = find_matching_bracket(text, end_label + 1, "(", ")")
    if end_url is None:
        return None
    label = text[start + 1 : end_label]
    url = text[end_label + 2 : end_url].strip()
    if url.startswith("<") and url.endswith(">"):
        url = url[1:-1].strip()
    if url:
        url = url.split()[0]
    return label, url, end_url + 1


def parse_autolink(text: str, start: int) -> Optional[tuple[str, int]]:
    if text[start] != "<":
        return None
    end = text.find(">", start + 1)
    if end == -1:
        return None
    content = text[start + 1 : end].strip()
    if not content:
        return None
    if content.startswith(("http://", "https://", "mailto:")):
        return content, end + 1
    if "@" in content and " " not in content:
        return content, end + 1
    return None


def find_matching_bracket(text: str, start: int, open_char: str, close_char: str) -> Optional[int]:
    depth = 0
    i = start
    while i < len(text):
        ch = text[i]
        if ch == "\\":
            i += 2
            continue
        if ch == open_char:
            depth += 1
        elif ch == close_char:
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return None


def can_toggle_marker(text: str, start: int, marker: str) -> bool:
    if not marker.startswith("_"):
        return True
    prev_char = text[start - 1] if start > 0 else ""
    next_index = start + len(marker)
    next_char = text[next_index] if next_index < len(text) else ""
    if prev_char and next_char and is_word_char(prev_char) and is_word_char(next_char):
        return False
    return True


def is_word_char(ch: str) -> bool:
    return ch.isalnum() or ch == "_"


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_table(
    doc: Document,
    header: list[str],
    rows: list[list[str]],
    alignments: list[WD_ALIGN_PARAGRAPH],
    args: argparse.Namespace,
) -> None:
    col_count = max([len(header)] + [len(row) for row in rows] + [0])
    if col_count == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_layout_fixed(table)

    text_width_cm = 21.0 - args.margin_left_cm - args.margin_right_cm
    col_width = text_width_cm / col_count if col_count else text_width_cm
    for column in table.columns:
        column.width = Cm(col_width)

    set_table_cell_margins(table, top_cm=0.1, bottom_cm=0.1, left_cm=0.2, right_cm=0.2)

    def fill_row(row_index: int, data: list[str], bold: bool) -> None:
        for col_index in range(col_count):
            text = data[col_index] if col_index < len(data) else ""
            cell = table.cell(row_index, col_index)
            cell.text = ""
            cell.width = Cm(col_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
            if col_index < len(alignments):
                paragraph.alignment = alignments[col_index]
            add_text_with_breaks(paragraph, text, args, bold)

    fill_row(0, header, True)
    for row_idx, row in enumerate(rows, start=1):
        fill_row(row_idx, row, False)


def markdown_to_docx(md_text: str, doc: Document, args: argparse.Namespace) -> None:
    lines = md_text.splitlines()
    heading_re = re.compile(r"^(#{1,5})\s+(.*)$")
    ordered_re = re.compile(r"^\d+[.)、]\s+")
    bullet_re = re.compile(r"^[-+*]\s+")
    task_re = re.compile(r"^[-+*]\s+\[( |x|X)\]\s+")
    header_keys = ["份号", "密级", "紧急程度", "发文机关标志", "发文字号", "签发人"]
    header_re = re.compile(rf"^({'|'.join(map(re.escape, header_keys))})[:：]\s*(.*)$")
    recipient_re = re.compile(r"^主送(?:机关)?[:：]\s*(.*)$")
    attachment_re = re.compile(r"^附件[:：]\s*(.*)$")
    copy_re = re.compile(r"^抄送(?:机关)?[:：]\s*(.*)$")
    print_re = re.compile(r"^印发(?:机关|单位)?(?:及日期)?[:：]\s*(.*)$")
    signature_re = re.compile(r"^(落款|署名|发文机关署名)[:：]?\s*(.*)$")

    header_seen = False
    separator_inserted = False
    signature_mode = False
    in_code_block = False
    code_block_lines: list[str] = []
    paragraph_buffer: list[tuple[str, bool]] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(args.line_spacing_pt)
        for text, hard_break in paragraph_buffer:
            add_inline_runs(
                paragraph,
                text,
                args,
                bold=None,
                font_name=args.font,
                font_size=args.font_size,
            )
            if hard_break:
                run = paragraph.add_run("")
                apply_run_format(
                    run,
                    args,
                    bold=None,
                    italic=None,
                    strike=None,
                    font_name=args.font,
                    font_size=args.font_size,
                )
                run.add_break()
        paragraph_buffer.clear()

    def append_paragraph_segment(text: str, hard_break: bool) -> None:
        text = text.strip()
        if not text and not hard_break:
            return
        if not paragraph_buffer:
            paragraph_buffer.append((text, hard_break))
            return
        last_text, last_break = paragraph_buffer[-1]
        if last_break:
            paragraph_buffer.append((text, hard_break))
            return
        if text:
            if needs_space(last_text, text):
                last_text = f"{last_text} {text}"
            else:
                last_text = f"{last_text}{text}"
        paragraph_buffer[-1] = (last_text, hard_break)

    def append_paragraph_line(raw_line: str) -> None:
        for segment, inline_break in split_inline_breaks(raw_line):
            text, hard_break = extract_manual_break(segment)
            append_paragraph_segment(text, hard_break or inline_break)

    i = 0
    while i < len(lines):
        raw_line = lines[i]
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code_block:
                add_code_block(doc, code_block_lines, args)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue
        if in_code_block:
            code_block_lines.append(raw_line.rstrip("\n"))
            i += 1
            continue

        if stripped == "":
            flush_paragraph()
            if signature_mode:
                signature_mode = False
            i += 1
            continue

        header_match = header_re.match(stripped)
        if header_match:
            flush_paragraph()
            key, value = header_match.groups()
            value = value.strip()
            header_seen = True
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            if key == "发文机关标志":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(value)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                set_run_font(run, args.title_font, args.title_size, False, args.digit_font)
            elif key == "发文字号":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(value)
                set_run_font(run, args.font, args.heading_size, None, args.digit_font)
            elif key == "份号":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(value)
                set_run_font(run, args.digit_font, args.heading_size, None, args.digit_font)
            elif key == "签发人":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                label_run = paragraph.add_run("签发人")
                set_run_font(label_run, args.font, args.heading_size, None, args.digit_font)
                if value:
                    name_run = paragraph.add_run(f" {value}")
                    set_run_font(name_run, args.heading2_font, args.heading_size, None, args.digit_font)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(value)
                set_run_font(run, args.heading1_font, args.heading_size, None, args.digit_font)
            i += 1
            continue

        if (
            stripped == "---"
            or re.match(r"^-{3,}$", stripped)
            or re.match(r"^[*_]{3,}$", stripped)
        ):
            flush_paragraph()
            add_red_separator(doc)
            separator_inserted = True
            i += 1
            continue

        if (
            is_table_row(stripped)
            and i + 1 < len(lines)
            and is_table_separator_line(lines[i + 1].strip())
        ):
            flush_paragraph()
            header = split_table_row(stripped)
            alignments = parse_table_alignments(lines[i + 1].strip())
            rows: list[list[str]] = []
            i += 2
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line or not is_table_row(row_line) or is_table_separator_line(row_line):
                    break
                rows.append(split_table_row(row_line))
                i += 1
            add_table(doc, header, rows, alignments, args)
            continue

        heading = heading_re.match(stripped)
        if heading:
            flush_paragraph()
            prefix = heading.group(1)
            title = heading.group(2).strip()
            if header_seen and not separator_inserted and len(prefix) == 1:
                add_red_separator(doc)
                separator_inserted = True
            heading_level = min(len(prefix), 5)
            paragraph = doc.add_paragraph(style=f"Heading {heading_level}")
            add_text_with_breaks(
                paragraph,
                title,
                args,
                bold=None,
                font_name=None,
                font_size=None,
            )
            i += 1
            continue

        signature_match = signature_re.match(stripped)
        if signature_match:
            flush_paragraph()
            signature_mode = True
            value = signature_match.group(2).strip()
            if value:
                add_right_paragraph(doc, value, args)
            i += 1
            continue

        if signature_mode:
            add_right_paragraph(doc, stripped, args)
            i += 1
            continue

        recipient_match = recipient_re.match(stripped)
        if recipient_match:
            flush_paragraph()
            value = recipient_match.group(1).strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            add_text_with_breaks(paragraph, f"主送：{value}", args, bold=None)
            i += 1
            continue

        attachment_match = attachment_re.match(stripped)
        if attachment_match:
            flush_paragraph()
            value = attachment_match.group(1).strip()
            paragraph = doc.add_paragraph()
            add_text_with_breaks(
                paragraph,
                f"附件：{value}" if value else "附件：",
                args,
                bold=None,
            )
            i += 1
            continue

        copy_match = copy_re.match(stripped)
        if copy_match:
            flush_paragraph()
            value = copy_match.group(1).strip()
            paragraph = doc.add_paragraph(style="Imprint")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            add_text_with_breaks(
                paragraph,
                f"抄送：{value}",
                args,
                bold=None,
                font_size=14,
            )
            i += 1
            continue

        print_match = print_re.match(stripped)
        if print_match:
            flush_paragraph()
            value = print_match.group(1).strip()
            paragraph = doc.add_paragraph(style="Imprint")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            add_text_with_breaks(
                paragraph,
                f"印发：{value}" if value else "印发：",
                args,
                bold=None,
                font_size=14,
            )
            i += 1
            continue

        task_match = task_re.match(stripped)
        if task_match:
            flush_paragraph()
            checked = task_match.group(1).lower() == "x"
            text = stripped[task_match.end() :].strip()
            paragraph = doc.add_paragraph(style="List Bullet")
            marker = "☑ " if checked else "☐ "
            add_text_with_breaks(paragraph, f"{marker}{text}", args, bold=None)
            i += 1
            continue

        if bullet_re.match(stripped):
            flush_paragraph()
            text = bullet_re.sub("", stripped, count=1)
            paragraph = doc.add_paragraph(style="List Bullet")
            add_text_with_breaks(paragraph, text.strip(), args, bold=None)
            i += 1
            continue

        if ordered_re.match(stripped):
            flush_paragraph()
            text = ordered_re.sub("", stripped, count=1)
            paragraph = doc.add_paragraph(style="List Number")
            add_text_with_breaks(paragraph, text.strip(), args, bold=None)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_text = stripped.lstrip(">").strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            add_text_with_breaks(paragraph, quote_text, args, bold=None)
            i += 1
            continue

        append_paragraph_line(raw_line)
        i += 1

    flush_paragraph()


def run_pandoc(input_path: Path, output_path: Path, reference_doc: Path) -> None:
    try:
        subprocess.run(
            [
                "pandoc",
                str(input_path),
                "-o",
                str(output_path),
                "--reference-doc",
                str(reference_doc),
            ],
            check=True,
        )
    except FileNotFoundError as exc:
        raise PandocNotFoundError(
            "pandoc not found in PATH. Run inside the Wunder container or install pandoc."
        ) from exc
    except subprocess.CalledProcessError as exc:
        raise PandocFailedError(f"pandoc failed with exit code {exc.returncode}") from exc


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"Input markdown not found: {input_path}", file=sys.stderr)
        return 1

    output_path = Path(args.output).resolve() if args.output else input_path.with_suffix(".docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    raw_text = input_path.read_text(encoding="utf-8-sig")
    try:
        normalized_text = apply_heading_numbering(raw_text)
    except ValueError as exc:
        print(str(exc), file=sys.stderr)
        return 1

    if args.use_pandoc:
        try:
            with TemporaryDirectory() as temp_dir:
                temp_dir_path = Path(temp_dir)
                normalized_md = temp_dir_path / "normalized.md"
                normalized_md.write_text(normalized_text, encoding="utf-8")

                if args.reference_doc:
                    reference_doc = Path(args.reference_doc).resolve()
                    if not reference_doc.exists():
                        print(f"Reference docx not found: {reference_doc}", file=sys.stderr)
                        return 1
                    run_pandoc(normalized_md, output_path, reference_doc)
                else:
                    reference_doc = temp_dir_path / "reference.docx"
                    build_reference_docx(reference_doc, args)
                    run_pandoc(normalized_md, output_path, reference_doc)
        except PandocNotFoundError:
            print("pandoc not found; run without --use-pandoc to use fallback.", file=sys.stderr)
            return 1
        except PandocFailedError as exc:
            print(str(exc), file=sys.stderr)
            return 1
    else:
        try:
            doc = create_base_document(args)
            markdown_to_docx(normalized_text, doc, args)
            doc.save(output_path)
        except PermissionError:
            print(f"Output file is locked: {output_path}", file=sys.stderr)
            return 1

    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
