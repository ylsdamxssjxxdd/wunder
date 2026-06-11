#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt


IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<target>[^)]+)\)")


def resolve_converter(explicit_path: str | None) -> Path | None:
    candidates: list[Path] = []

    if explicit_path:
        candidates.append(Path(explicit_path).expanduser())

    candidates.append(Path(__file__).resolve().with_name("convert_markdown_to_docx.py"))

    env_path = os.environ.get("GONGWEN_WRITER_CONVERTER")
    if env_path:
        candidates.append(Path(env_path).expanduser())

    desktop_candidate = (
        Path.home()
        / "Desktop"
        / "wunder"
        / "config"
        / "skills"
        / "公文写作"
        / "scripts"
        / "convert_markdown_to_docx.py"
    )
    candidates.append(desktop_candidate)

    for candidate in candidates:
        if candidate.is_file():
            return candidate.resolve()
    return None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Call a local Markdown-to-DOCX converter and postprocess embedded images.",
        epilog=(
            "Unknown options are forwarded to convert_markdown_to_docx.py. "
            "For image-heavy documents, pass --use-pandoc explicitly or let this wrapper "
            "add it automatically when local Markdown images are detected."
        ),
    )
    parser.add_argument("input_md", help="Input Markdown file")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument(
        "--converter",
        help="Explicit path to convert_markdown_to_docx.py",
    )
    args, extra_args = parser.parse_known_args()
    args.extra_args = [item for item in extra_args if item != "--"]
    return args


def collect_markdown_images(markdown_path: Path) -> list[tuple[str, Path]]:
    images: list[tuple[str, Path]] = []
    text = markdown_path.read_text(encoding="utf-8-sig")
    for match in IMAGE_RE.finditer(text):
        alt_text = match.group("alt").strip()
        relative_path = split_markdown_image_target(match.group("target"))
        if not relative_path or is_remote_image(relative_path):
            continue
        image_path = (markdown_path.parent / relative_path).resolve()
        images.append((alt_text, image_path))
    return images


def split_markdown_image_target(target: str) -> str:
    target = target.strip()
    if not target:
        return ""
    if target.startswith("<"):
        end = target.find(">")
        if end != -1:
            return target[1:end].strip()
    return target.split()[0]


def is_remote_image(path: str) -> bool:
    lowered = path.lower()
    return lowered.startswith(("http://", "https://", "data:", "file:"))


def has_converter_flag(args: list[str], flag: str) -> bool:
    return flag in args or any(item.startswith(f"{flag}=") for item in args)


def is_image_placeholder(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if stripped.startswith("图片：") or stripped.startswith("圖片："):
        return True
    if "../assets/" in stripped or ".png" in stripped or ".jpg" in stripped or ".jpeg" in stripped:
        return True
    return False


def configure_picture_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_caption_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(12)
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def embed_local_images_into_docx(markdown_path: Path, docx_path: Path) -> None:
    images = collect_markdown_images(markdown_path)
    if not images:
        return

    doc = Document(docx_path)
    image_index = 0

    for paragraph in doc.paragraphs:
        if image_index >= len(images):
            break
        if not is_image_placeholder(paragraph.text):
            continue

        alt_text, image_path = images[image_index]
        image_index += 1
        if not image_path.is_file():
            continue

        picture_paragraph = paragraph.insert_paragraph_before()
        configure_picture_paragraph(picture_paragraph)
        picture_run = picture_paragraph.add_run()
        picture_run.add_picture(str(image_path), width=Cm(12.8))

        paragraph.text = alt_text or image_path.name
        configure_caption_paragraph(paragraph)

    doc.save(docx_path)


def main() -> int:
    args = parse_args()

    input_md = Path(args.input_md).expanduser().resolve()
    if not input_md.is_file():
        print(f"Input file not found: {input_md}", file=sys.stderr)
        return 1

    output_docx = Path(args.output).expanduser().resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    converter = resolve_converter(args.converter)
    if converter is None:
        print(
            "No usable convert_markdown_to_docx.py found. "
            "Use --converter or set GONGWEN_WRITER_CONVERTER.",
            file=sys.stderr,
        )
        return 1

    extra_args = list(args.extra_args)
    markdown_images = collect_markdown_images(input_md)
    use_pandoc = has_converter_flag(extra_args, "--use-pandoc")
    if markdown_images and not use_pandoc:
        extra_args.append("--use-pandoc")
        use_pandoc = True

    command = [
        sys.executable,
        str(converter),
        str(input_md),
        "--output",
        str(output_docx),
        *extra_args,
    ]

    completed = subprocess.run(command, check=False)
    if completed.returncode == 0 and output_docx.is_file() and not use_pandoc:
        embed_local_images_into_docx(input_md, output_docx)
    return completed.returncode


if __name__ == "__main__":
    raise SystemExit(main())
