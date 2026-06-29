#!/usr/bin/env python3
"""Proofread DOCX and produce format-preserving annotated revisions."""

from __future__ import annotations

import argparse
from copy import deepcopy
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.run import Run
except Exception as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Please install it first: pip install python-docx\n"
        f"detail: {exc}"
    )


EXPECTED = {
    "page_width_cm": 21.0,
    "page_height_cm": 29.7,
    "margin_top_cm": 3.7,
    "margin_bottom_cm": 3.5,
    "margin_left_cm": 2.8,
    "margin_right_cm": 2.6,
    "title_font_keywords": ("方正小标宋", "小标宋"),
    "title_size_pt": 22.0,
    "heading1_font_keywords": ("黑体",),
    "heading2_font_keywords": ("楷体",),
    "heading3_font_keywords": ("仿宋",),
    "body_font_keywords": ("仿宋",),
    "body_size_pt": 16.0,
    "line_spacing_pt": 28.9,
    "first_line_indent_pt": 32.0,
}

TOLERANCE = {
    "page_cm": 0.2,
    "margin_cm": 0.2,
    "size_pt": 1.0,
    "line_spacing_pt": 1.2,
    "indent_pt": 4.0,
}

COMMON_TYPO_MAP = {
    "部暑": "部署",
    "布署": "部署",
    "必需": "必须",
    "做为": "作为",
    "以经": "已经",
    "即然": "既然",
    "按排": "安排",
    "决对": "绝对",
    "在职工做": "在职工作",
    "通迅": "通讯",
    "凭添": "平添",
    "渡过难关": "度过难关",
    "再接再励": "再接再厉",
    "一愁莫展": "一筹莫展",
    "默守成规": "墨守成规",
    "相形见拙": "相形见绌",
    "随声附合": "随声附和",
    "黙契": "默契",
    "按步就班": "按部就班",
    "安份守己": "安分守己",
    "重蹈复辙": "重蹈覆辙",
    "既往不究": "既往不咎",
    "冒然": "贸然",
    "针贬": "针砭",
    "精减": "精简",
    "幅射": "辐射",
}

DOUBLE_CHAR_TYPOS = {
    "的的": "的",
    "了了": "了",
    "在在": "在",
    "和和": "和",
    "是是": "是",
}

HEADING_PATTERNS = (
    ("heading1", re.compile(r"^[一二三四五六七八九十百千]+、")),
    ("heading2", re.compile(r"^（[一二三四五六七八九十百千]+）")),
    ("heading3", re.compile(r"^[0-9]+[\\.．]")),
    ("heading4", re.compile(r"^（[0-9]+）")),
)

STRUCTURE_PREFIXES = ("主送：", "附件：", "落款：", "抄送：", "印发：", "签发人：", "发文字号：")

DEFAULT_REVIEW_AUTHOR = "Wunder 文稿校对"


@dataclass
class Issue:
    severity: str
    category: str
    rule: str
    location: str
    snippet: str
    expected: str
    actual: str
    suggestion: str


@dataclass
class TextBlock:
    block_id: str
    kind: str
    location: str
    text: str
    normalized_text: str
    style: str
    char_count: int


@dataclass
class TextEdit:
    block_id: str
    before: str
    after: str
    reason: str
    severity: str
    category: str
    occurrence: int = 1


@dataclass
class AppliedChange:
    index: int
    block_id: str
    location: str
    before: str
    after: str
    reason: str
    severity: str
    category: str
    occurrence: int
    status: str
    message: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Proofread DOCX document formatting and typos.")
    parser.add_argument("docx", help="Input .docx path")
    parser.add_argument("--output-json", dest="output_json", help="Write result JSON to file")
    parser.add_argument("--output-md", dest="output_md", help="Write markdown report to file")
    parser.add_argument(
        "--output-blocks-json",
        dest="output_blocks_json",
        help="Write model-editable text blocks with stable block_id values.",
    )
    parser.add_argument(
        "--extract-blocks-only",
        action="store_true",
        help="Only extract DOCX text blocks; skip format/typo scoring.",
    )
    parser.add_argument(
        "--apply-edits",
        dest="apply_edits",
        help="Apply a JSON edit manifest to the original DOCX in place-preserving mode.",
    )
    parser.add_argument(
        "--output-docx",
        dest="output_docx",
        help="Write annotated revised DOCX. Defaults to '<input>-标注修订版.docx' when --apply-edits is used.",
    )
    parser.add_argument(
        "--output-clean-docx",
        dest="output_clean_docx",
        help="Optionally write a clean revised DOCX without highlight, comments, or change-list appendix.",
    )
    parser.add_argument(
        "--output-changes-json",
        dest="output_changes_json",
        help="Write structured edit-application results to JSON.",
    )
    parser.add_argument(
        "--output-changes-md",
        dest="output_changes_md",
        help="Write a Markdown change list for the revised document.",
    )
    parser.add_argument(
        "--review-author",
        default=DEFAULT_REVIEW_AUTHOR,
        help=f"Author name for Word comments (default: {DEFAULT_REVIEW_AUTHOR}).",
    )
    parser.add_argument(
        "--append-change-list",
        dest="append_change_list",
        action="store_true",
        help="Append an internal change-list table to the annotated DOCX. Off by default.",
    )
    parser.add_argument(
        "--max-findings",
        type=int,
        default=200,
        help="Maximum findings returned in issues/typo list (default: 200)",
    )
    return parser.parse_args()


def is_close(actual: Optional[float], target: float, tolerance: float) -> bool:
    if actual is None:
        return False
    return abs(actual - target) <= tolerance


def to_pt(value) -> Optional[float]:
    if value is None:
        return None
    if hasattr(value, "pt"):
        try:
            return float(value.pt)
        except Exception:
            return None
    try:
        return float(value)
    except Exception:
        return None


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def clip_text(text: str, limit: int = 48) -> str:
    text = normalize_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def paragraph_text_items(document) -> list[tuple[int, str, object]]:
    items = []
    for idx, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if text:
            items.append((idx, text, paragraph))
    return items


def iter_document_text_paragraphs(document) -> list[tuple[str, str, str, object]]:
    """Return stable text-bearing paragraph anchors for model edit manifests."""
    items: list[tuple[str, str, str, object]] = []
    for idx, paragraph in enumerate(document.paragraphs, start=1):
        items.append((f"p{idx:04d}", "paragraph", f"正文第{idx}段", paragraph))

    for table_idx, table in enumerate(document.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            for cell_idx, cell in enumerate(row.cells, start=1):
                for paragraph_idx, paragraph in enumerate(cell.paragraphs, start=1):
                    block_id = f"t{table_idx:03d}r{row_idx:03d}c{cell_idx:03d}p{paragraph_idx:03d}"
                    location = f"表{table_idx}第{row_idx}行第{cell_idx}列第{paragraph_idx}段"
                    items.append((block_id, "table_cell", location, paragraph))

    for section_idx, section in enumerate(document.sections, start=1):
        header = section.header
        for paragraph_idx, paragraph in enumerate(header.paragraphs, start=1):
            block_id = f"s{section_idx:03d}h{paragraph_idx:03d}"
            location = f"第{section_idx}节页眉第{paragraph_idx}段"
            items.append((block_id, "header", location, paragraph))
        footer = section.footer
        for paragraph_idx, paragraph in enumerate(footer.paragraphs, start=1):
            block_id = f"s{section_idx:03d}f{paragraph_idx:03d}"
            location = f"第{section_idx}节页脚第{paragraph_idx}段"
            items.append((block_id, "footer", location, paragraph))

    return items


def block_style_name(paragraph) -> str:
    style = getattr(paragraph, "style", None)
    name = getattr(style, "name", None)
    return str(name or "").strip()


def build_text_blocks(document) -> list[TextBlock]:
    blocks = []
    for block_id, kind, location, paragraph in iter_document_text_paragraphs(document):
        text = paragraph.text or ""
        normalized = normalize_text(text)
        if not normalized:
            continue
        blocks.append(
            TextBlock(
                block_id=block_id,
                kind=kind,
                location=location,
                text=text,
                normalized_text=normalized,
                style=block_style_name(paragraph),
                char_count=len(text),
            )
        )
    return blocks


def extract_blocks_payload(path: Path) -> dict:
    document = Document(str(path))
    blocks = build_text_blocks(document)
    return {
        "ok": True,
        "document": {
            "path": str(path),
            "block_count": len(blocks),
        },
        "edit_manifest_schema": {
            "edits": [
                {
                    "block_id": "p0001",
                    "before": "原文中的精确片段",
                    "after": "修订后的文本",
                    "reason": "修改原因，写给用户审阅",
                    "severity": "low|medium|high",
                    "category": "错别字|术语|标点|语病|格式|其他",
                    "occurrence": 1,
                }
            ]
        },
        "blocks": [asdict(block) for block in blocks],
    }


def paragraph_level(text: str) -> Optional[str]:
    for level, pattern in HEADING_PATTERNS:
        if pattern.search(text):
            return level
    return None


def extract_run_font_name(run, *, prefer_east_asia: bool = True) -> Optional[str]:
    east_asia_names = []
    fallback_names = []
    r_pr = getattr(run._element, "rPr", None)
    if r_pr is not None and getattr(r_pr, "rFonts", None) is not None:
        for key in ("eastAsia",):
            value = r_pr.rFonts.get(qn(f"w:{key}"))
            if value:
                east_asia_names.append(value)
        for key in ("ascii", "hAnsi", "cs"):
            value = r_pr.rFonts.get(qn(f"w:{key}"))
            if value:
                fallback_names.append(value)
    if run.font and run.font.name:
        fallback_names.append(run.font.name)

    candidates = east_asia_names if prefer_east_asia else [*fallback_names, *east_asia_names]
    for name in candidates:
        cleaned = str(name).strip()
        if cleaned:
            return cleaned
    return None


def style_font_name(style, *, prefer_east_asia: bool = True) -> Optional[str]:
    if style is None:
        return None
    element = getattr(style, "element", None)
    r_pr = getattr(element, "rPr", None)
    if r_pr is not None and getattr(r_pr, "rFonts", None) is not None:
        east_asia = r_pr.rFonts.get(qn("w:eastAsia"))
        if prefer_east_asia and east_asia:
            return str(east_asia).strip()
        if prefer_east_asia:
            return None
        for key in ("ascii", "hAnsi", "cs"):
            value = r_pr.rFonts.get(qn(f"w:{key}"))
            if value:
                return str(value).strip()
        if east_asia:
            return str(east_asia).strip()
    font = getattr(style, "font", None)
    name = getattr(font, "name", None)
    return str(name).strip() if name else None


def dominant_font(paragraph, *, prefer_east_asia: bool = True) -> Optional[str]:
    weights = {}
    for run in paragraph.runs:
        text_len = len(run.text.strip())
        if text_len == 0:
            continue
        font_name = extract_run_font_name(run, prefer_east_asia=prefer_east_asia)
        if not font_name:
            continue
        weights[font_name] = weights.get(font_name, 0) + text_len
    if not weights:
        return style_font_name(getattr(paragraph, "style", None), prefer_east_asia=prefer_east_asia)
    return max(weights.items(), key=lambda item: item[1])[0]


def dominant_size_pt(paragraph) -> Optional[float]:
    weights = {}
    for run in paragraph.runs:
        text_len = len(run.text.strip())
        if text_len == 0:
            continue
        size = to_pt(run.font.size)
        if size is None:
            continue
        rounded = round(size, 1)
        weights[rounded] = weights.get(rounded, 0) + text_len
    if weights:
        return max(weights.items(), key=lambda item: item[1])[0]
    style = getattr(paragraph, "style", None)
    if style is not None and style.font is not None:
        return to_pt(style.font.size)
    return None


def paragraph_line_spacing_pt(paragraph) -> Optional[float]:
    value = to_pt(paragraph.paragraph_format.line_spacing)
    if value is not None:
        return value
    style = getattr(paragraph, "style", None)
    if style is not None and style.paragraph_format is not None:
        return to_pt(style.paragraph_format.line_spacing)
    return None


def paragraph_first_indent_pt(paragraph) -> Optional[float]:
    value = to_pt(paragraph.paragraph_format.first_line_indent)
    if value is not None:
        return value
    style = getattr(paragraph, "style", None)
    if style is not None and style.paragraph_format is not None:
        return to_pt(style.paragraph_format.first_line_indent)
    return None


def font_matches(font_name: Optional[str], keywords: Iterable[str]) -> bool:
    if not font_name:
        return False
    lowered = font_name.lower()
    return any(keyword.lower() in lowered for keyword in keywords)


def add_issue(
    issues: list[Issue],
    severity: str,
    category: str,
    rule: str,
    location: str,
    snippet: str,
    expected: str,
    actual: str,
    suggestion: str,
) -> None:
    issues.append(
        Issue(
            severity=severity,
            category=category,
            rule=rule,
            location=location,
            snippet=snippet,
            expected=expected,
            actual=actual,
            suggestion=suggestion,
        )
    )


def locate_title(items: list[tuple[int, str, object]]) -> Optional[int]:
    if not items:
        return None
    for index, (_, text, _) in enumerate(items):
        if text.startswith(STRUCTURE_PREFIXES):
            continue
        return index
    return 0


def check_layout(document, issues: list[Issue]) -> None:
    if not document.sections:
        return
    section = document.sections[0]
    layout_checks = [
        ("页面宽度", section.page_width.cm, EXPECTED["page_width_cm"], TOLERANCE["page_cm"], "A4 宽度 21.0cm"),
        ("页面高度", section.page_height.cm, EXPECTED["page_height_cm"], TOLERANCE["page_cm"], "A4 高度 29.7cm"),
        ("上边距", section.top_margin.cm, EXPECTED["margin_top_cm"], TOLERANCE["margin_cm"], "3.7cm"),
        ("下边距", section.bottom_margin.cm, EXPECTED["margin_bottom_cm"], TOLERANCE["margin_cm"], "3.5cm"),
        ("左边距", section.left_margin.cm, EXPECTED["margin_left_cm"], TOLERANCE["margin_cm"], "2.8cm"),
        ("右边距", section.right_margin.cm, EXPECTED["margin_right_cm"], TOLERANCE["margin_cm"], "2.6cm"),
    ]
    for rule, actual, target, tolerance, expected_text in layout_checks:
        if not is_close(actual, target, tolerance):
            add_issue(
                issues,
                severity="high",
                category="格式",
                rule=rule,
                location="页面设置",
                snippet="",
                expected=expected_text,
                actual=f"{actual:.2f}cm",
                suggestion=f"将{rule}调整为 {target:.1f}cm。",
            )


def check_title(items: list[tuple[int, str, object]], title_index: int, issues: list[Issue]) -> None:
    paragraph_no, text, paragraph = items[title_index]
    location = f"第{paragraph_no}段"
    align = paragraph.alignment
    if align != WD_ALIGN_PARAGRAPH.CENTER:
        add_issue(
            issues,
            severity="high",
            category="格式",
            rule="标题对齐",
            location=location,
            snippet=clip_text(text),
            expected="居中对齐",
            actual=f"{align}",
            suggestion="将标题段落设置为居中对齐。",
        )
    font_name = dominant_font(paragraph)
    if font_name is not None and not font_matches(font_name, EXPECTED["title_font_keywords"]):
        add_issue(
            issues,
            severity="high",
            category="格式",
            rule="标题字体",
            location=location,
            snippet=clip_text(text),
            expected="方正小标宋简体（或小标宋）",
            actual=font_name or "未检测到",
            suggestion="将标题字体调整为方正小标宋简体。",
        )
    size_pt = dominant_size_pt(paragraph)
    if not is_close(size_pt, EXPECTED["title_size_pt"], TOLERANCE["size_pt"]):
        add_issue(
            issues,
            severity="medium",
            category="格式",
            rule="标题字号",
            location=location,
            snippet=clip_text(text),
            expected="2号（约 22pt）",
            actual="未检测到" if size_pt is None else f"{size_pt:.1f}pt",
            suggestion="将标题字号调整为 2号（22pt 左右）。",
        )


def check_body_and_headings(
    items: list[tuple[int, str, object]],
    title_index: int,
    issues: list[Issue],
) -> None:
    for idx, (paragraph_no, text, paragraph) in enumerate(items):
        if idx == title_index:
            continue
        location = f"第{paragraph_no}段"
        if text.startswith(STRUCTURE_PREFIXES):
            continue
        level = paragraph_level(text)
        font_name = dominant_font(paragraph)
        size_pt = dominant_size_pt(paragraph)

        if level == "heading1" and font_name is not None and not font_matches(font_name, EXPECTED["heading1_font_keywords"]):
            add_issue(
                issues,
                "medium",
                "格式",
                "一级标题字体",
                location,
                clip_text(text),
                "黑体三号",
                font_name or "未检测到",
                "将一级标题字体设为黑体。",
            )
        elif level == "heading2" and font_name is not None and not font_matches(font_name, EXPECTED["heading2_font_keywords"]):
            add_issue(
                issues,
                "medium",
                "格式",
                "二级标题字体",
                location,
                clip_text(text),
                "楷体_GB2312 三号",
                font_name or "未检测到",
                "将二级标题字体设为楷体_GB2312。",
            )
        elif level in {"heading3", "heading4"} and font_name is not None and not font_matches(font_name, EXPECTED["heading3_font_keywords"]):
            add_issue(
                issues,
                "low",
                "格式",
                "三级/四级标题字体",
                location,
                clip_text(text),
                "仿宋_GB2312 三号",
                font_name or "未检测到",
                "将三级/四级标题字体设为仿宋_GB2312。",
            )
        elif level is None:
            if font_name is not None and not font_matches(font_name, EXPECTED["body_font_keywords"]):
                add_issue(
                    issues,
                    "low",
                    "格式",
                    "正文字体",
                    location,
                    clip_text(text),
                    "仿宋_GB2312 三号",
                    font_name or "未检测到",
                    "将正文字体调整为仿宋_GB2312。",
                )
            if not is_close(size_pt, EXPECTED["body_size_pt"], TOLERANCE["size_pt"]):
                add_issue(
                    issues,
                    "low",
                    "格式",
                    "正文字号",
                    location,
                    clip_text(text),
                    "三号（约 16pt）",
                    "未检测到" if size_pt is None else f"{size_pt:.1f}pt",
                    "将正文字号调整为三号（16pt 左右）。",
                )
            line_spacing = paragraph_line_spacing_pt(paragraph)
            if not is_close(line_spacing, EXPECTED["line_spacing_pt"], TOLERANCE["line_spacing_pt"]):
                add_issue(
                    issues,
                    "medium",
                    "格式",
                    "正文行距",
                    location,
                    clip_text(text),
                    "固定值 28.9pt",
                    "未检测到" if line_spacing is None else f"{line_spacing:.1f}pt",
                    "将正文行距设置为固定值 28.9pt。",
                )
            first_indent = paragraph_first_indent_pt(paragraph)
            if not is_close(first_indent, EXPECTED["first_line_indent_pt"], TOLERANCE["indent_pt"]):
                add_issue(
                    issues,
                    "medium",
                    "格式",
                    "正文首行缩进",
                    location,
                    clip_text(text),
                    "首行缩进 2 字（约 32pt）",
                    "未检测到" if first_indent is None else f"{first_indent:.1f}pt",
                    "将正文首行缩进设置为 2 字。",
                )


def detect_typos(items: list[tuple[int, str, object]]) -> list[Issue]:
    findings: list[Issue] = []
    for paragraph_no, text, _ in items:
        location = f"第{paragraph_no}段"
        for wrong, correct in COMMON_TYPO_MAP.items():
            start = 0
            while True:
                idx = text.find(wrong, start)
                if idx < 0:
                    break
                snippet = clip_text(text[max(0, idx - 10) : idx + len(wrong) + 10])
                add_issue(
                    findings,
                    "medium",
                    "错别字",
                    "常见误写词",
                    location,
                    snippet,
                    f"建议使用“{correct}”",
                    f"检测到“{wrong}”",
                    f"将“{wrong}”改为“{correct}”。",
                )
                start = idx + len(wrong)

        for wrong, correct in DOUBLE_CHAR_TYPOS.items():
            if wrong in text:
                add_issue(
                    findings,
                    "low",
                    "错别字",
                    "重复字",
                    location,
                    clip_text(text),
                    f"建议改为“{correct}”",
                    f"检测到“{wrong}”",
                    f"删除多余字符，建议改为“{correct}”。",
                )

        for match in re.finditer(r"([，。；：、！？])\1+", text):
            punct = match.group(0)
            add_issue(
                findings,
                "low",
                "文本",
                "重复标点",
                location,
                clip_text(text),
                "单个标点",
                f"检测到“{punct}”",
                "删除重复标点，保留一个。",
            )

        for match in re.finditer(r"([\\u4e00-\\u9fff])\\1\\1+", text):
            repeated = match.group(0)
            add_issue(
                findings,
                "low",
                "错别字",
                "疑似重复字",
                location,
                clip_text(text),
                "避免连续 3 个及以上相同汉字",
                f"检测到“{repeated}”",
                "检查是否为输入重复，必要时删除多余字符。",
            )

        for match in re.finditer(r"([\\u4e00-\\u9fff])[,:;]", text):
            punct = match.group(0)[-1]
            add_issue(
                findings,
                "low",
                "文本",
                "半角标点",
                location,
                clip_text(text),
                "中文语境建议使用全角标点（，：；）",
                f"检测到半角“{punct}”",
                "将半角标点替换为中文全角标点。",
            )
    return findings


def load_edit_manifest(path: Path) -> list[TextEdit]:
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    raw_edits = payload.get("edits", payload) if isinstance(payload, dict) else payload
    if not isinstance(raw_edits, list):
        raise ValueError("edit manifest must be a JSON array or an object with an edits array")

    edits: list[TextEdit] = []
    for index, item in enumerate(raw_edits, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"edit #{index} must be an object")
        block_id = str(item.get("block_id") or "").strip()
        before = str(item.get("before") or "")
        after = str(item.get("after") or "")
        reason = str(item.get("reason") or "").strip()
        severity = str(item.get("severity") or "medium").strip() or "medium"
        category = str(item.get("category") or "文本").strip() or "文本"
        occurrence = item.get("occurrence", 1)
        try:
            occurrence_int = int(occurrence)
        except Exception:
            occurrence_int = 1
        if not block_id:
            raise ValueError(f"edit #{index} missing block_id")
        if before == "":
            raise ValueError(f"edit #{index} missing before text")
        edits.append(
            TextEdit(
                block_id=block_id,
                before=before,
                after=after,
                reason=reason,
                severity=severity,
                category=category,
                occurrence=max(1, occurrence_int),
            )
        )
    return edits


def paragraph_map(document) -> dict[str, tuple[str, str, object]]:
    return {
        block_id: (kind, location, paragraph)
        for block_id, kind, location, paragraph in iter_document_text_paragraphs(document)
    }


def run_text(run) -> str:
    return run.text or ""


def clone_run_element(paragraph, source_run, text: str, highlight: bool) -> Run:
    new_r = OxmlElement("w:r")
    r_pr = getattr(source_run._r, "rPr", None)
    if r_pr is not None:
        new_r.append(deepcopy(r_pr))
    new_run = Run(new_r, paragraph)
    new_run.text = text
    if highlight:
        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_run


def insert_run_after(anchor_run, new_run: Run) -> None:
    anchor_run._r.addnext(new_run._r)


def remove_run(run) -> None:
    parent = run._r.getparent()
    if parent is not None:
        parent.remove(run._r)


def find_nth(haystack: str, needle: str, occurrence: int) -> int:
    start = 0
    for _ in range(max(1, occurrence)):
        index = haystack.find(needle, start)
        if index < 0:
            return -1
        start = index + len(needle)
    return index


def run_spans(paragraph) -> list[tuple[int, int, object, str]]:
    spans = []
    offset = 0
    for run in paragraph.runs:
        text = run_text(run)
        end = offset + len(text)
        spans.append((offset, end, run, text))
        offset = end
    return spans


def split_replacement_piece(
    replacement: str,
    source_segments: list[tuple[object, str]],
    highlight: bool,
    paragraph,
) -> list[Run]:
    if not replacement:
        return []

    source_lengths = [len(text) for _, text in source_segments if text]
    total = sum(source_lengths)
    if total <= 0:
        first_source = source_segments[0][0] if source_segments else paragraph.add_run("")
        return [clone_run_element(paragraph, first_source, replacement, highlight)]

    pieces: list[Run] = []
    consumed = 0
    for idx, (source_run, source_text) in enumerate(source_segments):
        if not source_text:
            continue
        if idx == len(source_segments) - 1:
            chunk = replacement[consumed:]
        else:
            proportional = round(len(replacement) * len(source_text) / total)
            remaining_sources = len(source_segments) - idx - 1
            max_end = len(replacement) - remaining_sources
            end = min(max(consumed + proportional, consumed), max_end)
            chunk = replacement[consumed:end]
            consumed = end
        if chunk:
            pieces.append(clone_run_element(paragraph, source_run, chunk, highlight))
    if not pieces:
        pieces.append(clone_run_element(paragraph, source_segments[0][0], replacement, highlight))
    return pieces


def replace_text_in_paragraph(paragraph, before: str, after: str, occurrence: int, highlight: bool) -> list[Run]:
    full_text = "".join(run_text(run) for run in paragraph.runs)
    start = find_nth(full_text, before, occurrence)
    if start < 0:
        return []
    end = start + len(before)

    spans = run_spans(paragraph)
    touched = [
        (run_start, run_end, run, text)
        for run_start, run_end, run, text in spans
        if run_end > start and run_start < end
    ]
    if not touched:
        return []

    first_start, _, first_run, first_text = touched[0]
    _, last_end, last_run, last_text = touched[-1]
    prefix = first_text[: max(0, start - first_start)]
    suffix = last_text[len(last_text) - max(0, last_end - end) :]

    source_segments = []
    for run_start, run_end, run, text in touched:
        segment_start = max(start, run_start) - run_start
        segment_end = min(end, run_end) - run_start
        source_segments.append((run, text[segment_start:segment_end]))

    first_run.text = prefix
    anchor = first_run
    inserted_runs = split_replacement_piece(after, source_segments, highlight, paragraph)
    for new_run in inserted_runs:
        insert_run_after(anchor, new_run)
        anchor = new_run

    if last_run is first_run:
        if suffix:
            suffix_run = clone_run_element(paragraph, first_run, suffix, highlight=False)
            insert_run_after(anchor, suffix_run)
            if not inserted_runs:
                inserted_runs = [suffix_run]
    else:
        last_run.text = suffix
        for _, _, run, _ in touched[1:-1]:
            remove_run(run)
        if not inserted_runs:
            inserted_runs = [last_run] if suffix else [first_run]

    return inserted_runs if inserted_runs else [first_run]


def add_review_comment(document, runs: list[Run], edit: TextEdit, review_author: str) -> bool:
    if not runs:
        return False
    text = f"{edit.before} -> {edit.after}"
    if edit.reason:
        text = f"{text}\n原因：{edit.reason}"
    try:
        document.add_comment(runs, text=text, author=review_author, initials="W")
        return True
    except Exception:
        return False


def apply_edits_to_document(
    document,
    edits: list[TextEdit],
    *,
    highlight: bool,
    add_comments: bool,
    review_author: str,
) -> list[AppliedChange]:
    blocks = paragraph_map(document)
    results: list[AppliedChange] = []
    for index, edit in enumerate(edits, start=1):
        block = blocks.get(edit.block_id)
        if block is None:
            results.append(
                AppliedChange(
                    index=index,
                    block_id=edit.block_id,
                    location="",
                    before=edit.before,
                    after=edit.after,
                    reason=edit.reason,
                    severity=edit.severity,
                    category=edit.category,
                    occurrence=edit.occurrence,
                    status="skipped",
                    message="block_id not found",
                )
            )
            continue
        _, location, paragraph = block
        inserted_runs = replace_text_in_paragraph(
            paragraph,
            edit.before,
            edit.after,
            edit.occurrence,
            highlight=highlight,
        )
        if not inserted_runs:
            results.append(
                AppliedChange(
                    index=index,
                    block_id=edit.block_id,
                    location=location,
                    before=edit.before,
                    after=edit.after,
                    reason=edit.reason,
                    severity=edit.severity,
                    category=edit.category,
                    occurrence=edit.occurrence,
                    status="skipped",
                    message="before text not found in target block",
                )
            )
            continue
        if add_comments:
            add_review_comment(document, inserted_runs, edit, review_author)
        results.append(
            AppliedChange(
                index=index,
                block_id=edit.block_id,
                location=location,
                before=edit.before,
                after=edit.after,
                reason=edit.reason,
                severity=edit.severity,
                category=edit.category,
                occurrence=edit.occurrence,
                status="applied",
                message="applied",
            )
        )
    return results


def append_change_list(document, changes: list[AppliedChange]) -> None:
    applied = [change for change in changes if change.status == "applied"]
    skipped = [change for change in changes if change.status != "applied"]
    document.add_page_break()
    document.add_heading("文稿校对修改清单", level=1)
    document.add_paragraph(f"已应用修改：{len(applied)} 处；未应用：{len(skipped)} 处。")
    if applied:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["序号", "位置", "类型", "原文", "修订", "原因"]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        for change in applied:
            row = table.add_row().cells
            row[0].text = str(change.index)
            row[1].text = change.location
            row[2].text = change.category
            row[3].text = change.before
            row[4].text = change.after
            row[5].text = change.reason
    if skipped:
        document.add_paragraph("未应用修改")
        for change in skipped:
            document.add_paragraph(
                f"{change.index}. {change.block_id}：{change.message}；原文片段：{change.before}",
                style=None,
            )


def default_revised_docx_path(path: Path) -> Path:
    return path.with_name(f"{path.stem}-标注修订版{path.suffix}")


def default_clean_docx_path(path: Path) -> Path:
    return path.with_name(f"{path.stem}-清洁修订版{path.suffix}")


def changes_markdown(source_path: Path, output_docx: Path, changes: list[AppliedChange]) -> str:
    applied = [change for change in changes if change.status == "applied"]
    skipped = [change for change in changes if change.status != "applied"]
    lines = [
        "# 文稿校对修改清单",
        "",
        f"- 原文档：`{source_path}`",
        f"- 标注版：`{output_docx}`",
        f"- 已应用修改：{len(applied)} 处",
        f"- 未应用修改：{len(skipped)} 处",
        "",
        "## 已应用修改",
    ]
    if not applied:
        lines.append("- 无。")
    else:
        for change in applied:
            lines.append(
                f"- {change.index}. {change.location} [{change.category}/{change.severity}] "
                f"`{change.before}` -> `{change.after}`；{change.reason or '未填写原因'}"
            )
    if skipped:
        lines.extend(["", "## 未应用修改"])
        for change in skipped:
            lines.append(
                f"- {change.index}. `{change.block_id}`：{change.message}；原文片段 `{change.before}`"
            )
    lines.append("")
    return "\n".join(lines)


def apply_edit_manifest(
    source_path: Path,
    manifest_path: Path,
    output_docx: Path,
    *,
    output_clean_docx: Optional[Path],
    append_list: bool,
    review_author: str,
) -> dict:
    edits = load_edit_manifest(manifest_path)
    document = Document(str(source_path))
    changes = apply_edits_to_document(
        document,
        edits,
        highlight=True,
        add_comments=True,
        review_author=review_author,
    )
    if append_list:
        append_change_list(document, changes)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))

    clean_path_value = None
    if output_clean_docx is not None:
        clean_document = Document(str(source_path))
        apply_edits_to_document(
            clean_document,
            edits,
            highlight=False,
            add_comments=False,
            review_author=review_author,
        )
        output_clean_docx.parent.mkdir(parents=True, exist_ok=True)
        clean_document.save(str(output_clean_docx))
        clean_path_value = str(output_clean_docx)

    applied_count = sum(1 for change in changes if change.status == "applied")
    skipped_count = len(changes) - applied_count
    return {
        "ok": skipped_count == 0,
        "document": {
            "source_path": str(source_path),
            "annotated_docx_path": str(output_docx),
            "clean_docx_path": clean_path_value,
        },
        "summary": {
            "requested_edit_count": len(edits),
            "applied_edit_count": applied_count,
            "skipped_edit_count": skipped_count,
        },
        "changes": [asdict(change) for change in changes],
    }


def deduplicate_issues(issues: list[Issue]) -> list[Issue]:
    seen = set()
    output = []
    for issue in issues:
        key = (
            issue.category,
            issue.rule,
            issue.location,
            issue.expected,
            issue.actual,
            issue.snippet,
        )
        if key in seen:
            continue
        seen.add(key)
        output.append(issue)
    return output


def sort_issues(issues: list[Issue]) -> list[Issue]:
    priority = {"high": 0, "medium": 1, "low": 2}
    return sorted(
        issues,
        key=lambda item: (
            priority.get(item.severity, 9),
            item.category,
            item.rule,
            item.location,
        ),
    )


def evaluate_grade(issues: list[Issue]) -> tuple[int, str, bool]:
    weight = {"high": 10, "medium": 5, "low": 2}
    score = 100 - sum(weight.get(issue.severity, 2) for issue in issues)
    score = max(0, score)
    if score >= 90:
        grade = "A"
    elif score >= 75:
        grade = "B"
    else:
        grade = "C"
    can_publish = not any(issue.severity == "high" for issue in issues)
    return score, grade, can_publish


def markdown_report(
    doc_path: str,
    score: int,
    grade: str,
    can_publish: bool,
    format_issues: list[Issue],
    typo_issues: list[Issue],
) -> str:
    lines = [
        "# 文稿校对结果",
        "",
        "## 一、总体结论",
        f"- 合规等级：**{grade}**",
        f"- 评分：**{score}/100**",
        f"- 是否建议直接发文：**{'是' if can_publish else '否'}**",
        f"- 校对文档：`{doc_path}`",
        "",
        "## 二、格式问题",
    ]
    if not format_issues:
        lines.append("- 未检出明显格式问题。")
    else:
        for issue in format_issues:
            lines.append(
                f"- [{issue.severity}] {issue.location} {issue.rule}："
                f"{issue.actual}（建议：{issue.suggestion}）"
            )

    lines.extend(["", "## 三、错别字与文本问题"])
    if not typo_issues:
        lines.append("- 未检出明显错别字或文本问题。")
    else:
        for issue in typo_issues:
            lines.append(
                f"- [{issue.severity}] {issue.location} {issue.rule}："
                f"{issue.actual}（建议：{issue.suggestion}）"
            )

    lines.extend(["", "## 四、优先修复清单（Top 3）"])
    top3 = [*format_issues, *typo_issues][:3]
    if not top3:
        lines.append("- 无。")
    else:
        for issue in top3:
            lines.append(f"- {issue.location} {issue.rule}：{issue.suggestion}")

    lines.extend(
        [
            "",
            "## 五、复核建议",
            "- 修订后建议再次运行同一脚本进行二次校对。",
            "- 本报告用于初筛，正式发文前建议人工终审。",
            "",
        ]
    )
    return "\n".join(lines)


def analyze_docx(path: Path, max_findings: int) -> dict:
    document = Document(str(path))
    items = paragraph_text_items(document)
    issues: list[Issue] = []
    check_layout(document, issues)

    title_idx = locate_title(items)
    if title_idx is None:
        add_issue(
            issues,
            "high",
            "结构",
            "标题缺失",
            "文档整体",
            "",
            "存在明确公文标题",
            "未检测到标题段落",
            "在文档前部补充标题，并设置为小标宋二号居中。",
        )
    else:
        check_title(items, title_idx, issues)
        check_body_and_headings(items, title_idx, issues)

    typo_issues = detect_typos(items)
    all_issues = deduplicate_issues([*issues, *typo_issues])
    all_issues = sort_issues(all_issues)[: max(1, max_findings)]

    format_issues = [issue for issue in all_issues if issue.category == "格式" or issue.category == "结构"]
    typo_or_text_issues = [issue for issue in all_issues if issue.category in {"错别字", "文本"}]
    score, grade, can_publish = evaluate_grade(all_issues)

    md = markdown_report(
        doc_path=str(path),
        score=score,
        grade=grade,
        can_publish=can_publish,
        format_issues=format_issues,
        typo_issues=typo_or_text_issues,
    )
    severity_count = {"high": 0, "medium": 0, "low": 0}
    for issue in all_issues:
        severity_count[issue.severity] = severity_count.get(issue.severity, 0) + 1

    return {
        "ok": True,
        "document": {
            "path": str(path),
            "paragraph_count": len(document.paragraphs),
            "non_empty_paragraph_count": len(items),
        },
        "summary": {
            "score": score,
            "grade": grade,
            "can_publish_directly": can_publish,
            "issue_count": len(all_issues),
            "severity_count": severity_count,
            "format_issue_count": len(format_issues),
            "typo_issue_count": len(typo_or_text_issues),
        },
        "issues": [asdict(issue) for issue in all_issues],
        "standard_report_markdown": md,
    }


def write_text(path: str, content: str) -> None:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")


def write_json(path: str, payload: dict) -> None:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    args = parse_args()
    docx_path = Path(args.docx)

    if not docx_path.exists():
        print(
            json.dumps(
                {"ok": False, "error": f"文件不存在: {docx_path}"},
                ensure_ascii=False,
            )
        )
        return 2
    if docx_path.suffix.lower() != ".docx":
        print(
            json.dumps(
                {
                    "ok": False,
                    "error": "仅支持 .docx 文件，请先将 .doc 另存为 .docx 后再校对。",
                },
                ensure_ascii=False,
            )
        )
        return 2

    if args.extract_blocks_only or args.output_blocks_json:
        blocks_result = extract_blocks_payload(docx_path)
        if args.output_blocks_json:
            write_json(args.output_blocks_json, blocks_result)
        if args.extract_blocks_only:
            print(json.dumps(blocks_result, ensure_ascii=False, indent=2))
            return 0

    if args.apply_edits:
        manifest_path = Path(args.apply_edits)
        if not manifest_path.exists():
            print(
                json.dumps(
                    {"ok": False, "error": f"修改清单不存在: {manifest_path}"},
                    ensure_ascii=False,
                )
            )
            return 2
        output_docx = Path(args.output_docx) if args.output_docx else default_revised_docx_path(docx_path)
        output_clean_docx = Path(args.output_clean_docx) if args.output_clean_docx else None
        try:
            apply_result = apply_edit_manifest(
                docx_path,
                manifest_path,
                output_docx,
                output_clean_docx=output_clean_docx,
                append_list=args.append_change_list,
                review_author=args.review_author,
            )
        except Exception as exc:
            print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
            return 1
        if args.output_changes_json:
            write_json(args.output_changes_json, apply_result)
        if args.output_changes_md:
            write_text(args.output_changes_md, changes_markdown(docx_path, output_docx, [
                AppliedChange(**item) for item in apply_result["changes"]
            ]))
        print(json.dumps(apply_result, ensure_ascii=False, indent=2))
        return 0

    result = analyze_docx(docx_path, max_findings=args.max_findings)
    if args.output_json:
        write_json(args.output_json, result)
    if args.output_md:
        write_text(args.output_md, result["standard_report_markdown"])
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
