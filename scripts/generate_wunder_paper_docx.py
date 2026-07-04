#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the Wunder multi-agent paper DOCX with readable embedded figures.

This script is intentionally local and deterministic: it reads the UTF-8 paper
Markdown, redraws Mermaid blocks as PNG files with CJK fonts, converts the
Markdown through the local official-document converter, then replaces image
placeholders with real inline pictures.
"""

from __future__ import annotations

import math
import re
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[1]
DESKTOP_DIR = Path.home() / "Desktop"
OUTPUT_DIR = DESKTOP_DIR / "wunder_paper_docx"
ASSETS_DIR = OUTPUT_DIR / "assets"
DOCX_PATH = OUTPUT_DIR / "wunder_multi_agent_platform_paper.docx"
SOURCE_MD_PATH = OUTPUT_DIR / "paper_docx_source.md"
CONTACT_SHEET_PATH = ASSETS_DIR / "figures_contact_sheet.png"
PAPER_TITLE = "面向多用户场景的多智能体调控平台设计与实现"

FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
)

PALETTE = {
    "ink": "#111827",
    "muted": "#4B5563",
    "line": "#374151",
    "blue": "#EAF2FF",
    "blue_edge": "#4B79B8",
    "teal": "#E6F6F2",
    "teal_edge": "#2D8C78",
    "amber": "#FFF4D9",
    "amber_edge": "#B98718",
    "green": "#EAF7EA",
    "green_edge": "#4C9A58",
    "rose": "#FDEEEF",
    "rose_edge": "#C7505A",
    "gray": "#F5F7FA",
    "gray_edge": "#9AA5B1",
    "purple": "#F1EDFF",
    "purple_edge": "#7762B5",
}


@dataclass
class FigureBlock:
    index: int
    title: str
    code: str


class FontBook:
    def __init__(self) -> None:
        font_path = next((path for path in FONT_CANDIDATES if path.exists()), None)
        if font_path is None:
            self.path = None
        else:
            self.path = str(font_path)

    def get(self, size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
        if self.path:
            return ImageFont.truetype(self.path, size=size)
        return ImageFont.load_default()


FONT_BOOK = FontBook()


class Diagram:
    def __init__(self, width: int, height: int, title: str) -> None:
        self.width = width
        self.height = height
        self.image = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.title = title
        self.title_font = FONT_BOOK.get(40, bold=True)
        self.node_font = FONT_BOOK.get(28)
        self.small_font = FONT_BOOK.get(23)
        self.caption_font = FONT_BOOK.get(24)
        self.draw_title()

    def draw_title(self) -> None:
        if not self.title:
            return
        self.draw.text(
            (self.width // 2, 34),
            self.title,
            fill=PALETTE["ink"],
            font=self.title_font,
            anchor="mt",
        )
        self.draw.line((90, 92, self.width - 90, 92), fill="#D6DDE6", width=2)

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.image.save(path, "PNG", dpi=(300, 300), optimize=True)

    def text_size(self, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
        bbox = self.draw.multiline_textbbox((0, 0), text, font=font, spacing=6)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]

    def wrap_text(self, text: str, max_width: int, font: ImageFont.ImageFont) -> str:
        text = clean_label(text)
        lines: list[str] = []
        for raw_line in text.splitlines() or [""]:
            line = ""
            for ch in raw_line:
                trial = line + ch
                if line and self.text_size(trial, font)[0] > max_width:
                    lines.append(line)
                    line = ch
                else:
                    line = trial
            if line:
                lines.append(line)
        return "\n".join(lines)

    def box(
        self,
        xy: tuple[int, int, int, int],
        text: str,
        fill: str = "#F8FAFC",
        outline: str = "#607D8B",
        radius: int = 20,
        font: ImageFont.ImageFont | None = None,
        stroke_width: int = 3,
    ) -> None:
        font = font or self.node_font
        self.draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=stroke_width)
        x1, y1, x2, y2 = xy
        label = self.wrap_text(text, max(40, x2 - x1 - 28), font)
        self.draw.multiline_text(
            ((x1 + x2) // 2, (y1 + y2) // 2),
            label,
            fill=PALETTE["ink"],
            font=font,
            anchor="mm",
            align="center",
            spacing=7,
        )

    def diamond(
        self,
        xy: tuple[int, int, int, int],
        text: str,
        fill: str = "#FFF4D9",
        outline: str = "#B98718",
    ) -> None:
        x1, y1, x2, y2 = xy
        points = [((x1 + x2) // 2, y1), (x2, (y1 + y2) // 2), ((x1 + x2) // 2, y2), (x1, (y1 + y2) // 2)]
        self.draw.polygon(points, fill=fill, outline=outline)
        self.draw.line(points + [points[0]], fill=outline, width=3)
        label = self.wrap_text(text, max(40, x2 - x1 - 54), self.node_font)
        self.draw.multiline_text(
            ((x1 + x2) // 2, (y1 + y2) // 2),
            label,
            fill=PALETTE["ink"],
            font=self.node_font,
            anchor="mm",
            align="center",
            spacing=6,
        )

    def arrow(
        self,
        start: tuple[int, int],
        end: tuple[int, int],
        label: str = "",
        color: str = "#455A64",
        width: int = 4,
        dashed: bool = False,
    ) -> None:
        if dashed:
            self.dashed_line(start, end, fill=color, width=width)
        else:
            self.draw.line((*start, *end), fill=color, width=width)
        self.arrow_head(start, end, color, width)
        if label:
            mx = (start[0] + end[0]) // 2
            my = (start[1] + end[1]) // 2
            self.label((mx, my - 16), label)

    def elbow_arrow(
        self,
        points: list[tuple[int, int]],
        label: str = "",
        color: str = "#455A64",
        width: int = 4,
        dashed: bool = False,
    ) -> None:
        for first, second in zip(points, points[1:]):
            if dashed:
                self.dashed_line(first, second, fill=color, width=width)
            else:
                self.draw.line((*first, *second), fill=color, width=width)
        if len(points) >= 2:
            self.arrow_head(points[-2], points[-1], color, width)
        if label and len(points) >= 2:
            mid = points[len(points) // 2]
            self.label((mid[0], mid[1] - 16), label)

    def dashed_line(
        self,
        start: tuple[int, int],
        end: tuple[int, int],
        fill: str,
        width: int,
        dash: int = 18,
        gap: int = 10,
    ) -> None:
        x1, y1 = start
        x2, y2 = end
        length = math.hypot(x2 - x1, y2 - y1)
        if length == 0:
            return
        dx = (x2 - x1) / length
        dy = (y2 - y1) / length
        pos = 0.0
        while pos < length:
            segment_end = min(pos + dash, length)
            self.draw.line(
                (
                    int(x1 + dx * pos),
                    int(y1 + dy * pos),
                    int(x1 + dx * segment_end),
                    int(y1 + dy * segment_end),
                ),
                fill=fill,
                width=width,
            )
            pos += dash + gap

    def arrow_head(self, start: tuple[int, int], end: tuple[int, int], color: str, width: int) -> None:
        sx, sy = start
        ex, ey = end
        angle = math.atan2(ey - sy, ex - sx)
        size = 16 + width
        left = (ex - size * math.cos(angle - math.pi / 6), ey - size * math.sin(angle - math.pi / 6))
        right = (ex - size * math.cos(angle + math.pi / 6), ey - size * math.sin(angle + math.pi / 6))
        self.draw.polygon([(ex, ey), left, right], fill=color)

    def label(self, center: tuple[int, int], text: str) -> None:
        label = self.wrap_text(text, 120, self.small_font)
        w, h = self.text_size(label, self.small_font)
        x, y = center
        pad = 8
        self.draw.rounded_rectangle(
            (x - w // 2 - pad, y - h // 2 - pad, x + w // 2 + pad, y + h // 2 + pad),
            radius=10,
            fill="white",
            outline="#D1D5DB",
            width=1,
        )
        self.draw.multiline_text((x, y), label, fill=PALETTE["muted"], font=self.small_font, anchor="mm", align="center")


def clean_label(text: str) -> str:
    text = text.strip().strip('"')
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    return re.sub(r"\s+", " ", text).replace(" / ", "\n/ ")


def locate_paper_markdown() -> Path:
    candidates = sorted((REPO_ROOT / "docs").glob("*.md"))
    for path in candidates:
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            continue
        if PAPER_TITLE in text and "```mermaid" in text:
            return path
    raise FileNotFoundError("Could not locate the Wunder paper markdown under docs/.")


def extract_figures(markdown: str) -> list[FigureBlock]:
    pattern = re.compile(r"```mermaid\s*\n(.*?)\n```", re.S)
    figures: list[FigureBlock] = []
    for index, match in enumerate(pattern.finditer(markdown), start=1):
        before = markdown[: match.start()].splitlines()
        title = f"图 {index}"
        for line in reversed(before[-12:]):
            stripped = line.strip().strip("*")
            if stripped.startswith("图 "):
                title = stripped
                break
        figures.append(FigureBlock(index=index, title=title, code=match.group(1).strip()))
    return figures


def build_source_markdown(markdown: str, figures: list[FigureBlock]) -> str:
    iterator = iter(figures)

    def replace_block(_: re.Match[str]) -> str:
        figure = next(iterator)
        return f"![{figure.title}](assets/fig{figure.index}.png)"

    text = re.sub(r"```mermaid\s*\n(.*?)\n```", replace_block, markdown, flags=re.S)
    text = re.sub(r"\n---\n", "\n\n", text)
    text = text.replace("### 3.2 状态分离的工程意义", "## 3.2 状态分离的工程意义")
    return text


def parse_flow_nodes(code: str) -> dict[str, str]:
    nodes: dict[str, str] = {}
    for node_id, _, bracket_label, brace_label in re.findall(r"([A-Za-z][A-Za-z0-9_]*)\s*(\[[\"']([^\"']+)[\"']\]|\{\"([^\"]+)\"\})", code):
        nodes[node_id] = clean_label(bracket_label or brace_label)
    for graph_id, graph_label in re.findall(r"subgraph\s+([A-Za-z][A-Za-z0-9_]*)\s*\[\"([^\"]+)\"\]", code):
        nodes[graph_id] = clean_label(graph_label)
    return nodes


def parse_sequence(code: str) -> tuple[list[tuple[str, str]], list[tuple[str, str, str]]]:
    participants: list[tuple[str, str]] = []
    messages: list[tuple[str, str, str]] = []
    for raw in code.splitlines():
        line = raw.strip()
        part = re.match(r"participant\s+([A-Za-z][A-Za-z0-9_]*)\s+as\s+(.+)", line)
        if part:
            participants.append((part.group(1), clean_label(part.group(2))))
            continue
        msg = re.match(r"([A-Za-z][A-Za-z0-9_]*)\s*-+>>\s*([A-Za-z][A-Za-z0-9_]*):\s*(.+)", line)
        if msg:
            messages.append((msg.group(1), msg.group(2), clean_label(msg.group(3))))
    return participants, messages


def render_state_model(figure: FigureBlock, path: Path) -> None:
    labels = parse_flow_nodes(figure.code)
    diagram = Diagram(2000, 1280, figure.title)
    pos = {
        "User": (760, 135, 1240, 245),
        "Control": (760, 325, 1240, 455),
        "Task": (760, 535, 1240, 665),
        "Cognitive": (210, 780, 650, 930),
        "Collaboration": (780, 780, 1220, 930),
        "Durable": (1350, 780, 1790, 930),
        "Projection": (760, 1030, 1240, 1160),
    }
    colors = {
        "User": ("#F5F7FA", "#9AA5B1"),
        "Control": (PALETTE["rose"], PALETTE["rose_edge"]),
        "Task": (PALETTE["amber"], PALETTE["amber_edge"]),
        "Cognitive": (PALETTE["blue"], PALETTE["blue_edge"]),
        "Collaboration": (PALETTE["teal"], PALETTE["teal_edge"]),
        "Durable": (PALETTE["green"], PALETTE["green_edge"]),
        "Projection": (PALETTE["purple"], PALETTE["purple_edge"]),
    }
    for node, xy in pos.items():
        diagram.box(xy, labels.get(node, node), *colors[node])
    diagram.arrow((1000, 245), (1000, 325))
    diagram.arrow((1000, 455), (1000, 535))
    diagram.elbow_arrow([(900, 665), (560, 720), (430, 780)])
    diagram.arrow((1000, 665), (1000, 780))
    diagram.elbow_arrow([(1100, 665), (1510, 720), (1570, 780)])
    diagram.arrow((1220, 855), (1350, 855))
    diagram.elbow_arrow([(1570, 930), (1570, 995), (1240, 1095)])
    diagram.elbow_arrow([(760, 1095), (520, 995), (430, 930)], "不得反向修改", color="#C7505A", dashed=True)
    diagram.elbow_arrow([(840, 1030), (720, 910), (760, 610)], "不得反向修改", color="#C7505A", dashed=True)
    diagram.arrow((1000, 1160), (1000, 1215))
    diagram.draw.text((1000, 1226), "观察层只消费事件投影，不反写认知态与任务态", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mt")
    diagram.save(path)


def render_architecture(figure: FigureBlock, path: Path) -> None:
    labels = parse_flow_nodes(figure.code)
    diagram = Diagram(2600, 1460, figure.title)
    columns = [
        ("Access", ["Web", "Admin", "Desktop", "Cli", "Channel"], PALETTE["blue"], PALETTE["blue_edge"]),
        ("Control", ["Auth", "Permission", "Approval", "Gateway"], PALETTE["rose"], PALETTE["rose_edge"]),
        ("Runtime", ["Thread", "Orchestrator", "Mission", "Tools"], PALETTE["amber"], PALETTE["amber_edge"]),
        ("Projection", ["Events", "Replay", "Board"], PALETTE["purple"], PALETTE["purple_edge"]),
        ("Data", ["Storage", "Workspace", "Memory", "Knowledge"], PALETTE["green"], PALETTE["green_edge"]),
    ]
    x0 = 90
    col_w = 460
    gap = 40
    y0 = 145
    h = 1110
    centers: dict[str, tuple[int, int]] = {}
    for i, (group, members, fill, edge) in enumerate(columns):
        x = x0 + i * (col_w + gap)
        diagram.draw.rounded_rectangle((x, y0, x + col_w, y0 + h), radius=24, fill="#FBFCFE", outline=edge, width=4)
        diagram.draw.text((x + col_w // 2, y0 + 42), labels.get(group, group), fill=edge, font=FONT_BOOK.get(30, bold=True), anchor="mm")
        node_h = 122
        inner_gap = 34
        start_y = y0 + 105
        for j, node in enumerate(members):
            ny = start_y + j * (node_h + inner_gap)
            xy = (x + 45, ny, x + col_w - 45, ny + node_h)
            diagram.box(xy, labels.get(node, node), fill, edge, radius=18, font=FONT_BOOK.get(23))
            centers[node] = ((xy[0] + xy[2]) // 2, (xy[1] + xy[3]) // 2)
        centers[group] = (x + col_w // 2, y0 + h // 2)
    for left, right in [("Access", "Control"), ("Control", "Runtime"), ("Runtime", "Projection")]:
        lx, ly = centers[left]
        rx, ry = centers[right]
        diagram.arrow((lx + col_w // 2, ly), (rx - col_w // 2, ry))
    diagram.elbow_arrow([(centers["Runtime"][0], y0 + h + 25), (centers["Data"][0], y0 + h + 25), (centers["Data"][0], y0 + h)], "持久化")
    diagram.elbow_arrow([(centers["Control"][0], y0 + h + 75), (centers["Data"][0], y0 + h + 75), (centers["Data"][0], y0 + h)], "权限与配置")
    diagram.elbow_arrow([(centers["Projection"][0], y0 - 35), (centers["Access"][0], y0 - 35), (centers["Access"][0], y0)], "实时投影")
    diagram.draw.text((1300, 1335), "三种运行形态共享执行内核；治理、投影与存储按职责边界协作", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm")
    diagram.save(path)


def render_linear_path(figure: FigureBlock, path: Path) -> None:
    labels = parse_flow_nodes(figure.code)
    nodes = ["Goal", "Agent", "Ability", "Runtime", "Governance", "UI", "Observe"]
    diagram = Diagram(2600, 720, figure.title)
    box_w = 310
    box_h = 124
    x0 = 95
    y = 265
    fills = [PALETTE["blue"], PALETTE["teal"], PALETTE["green"], PALETTE["amber"], PALETTE["rose"], PALETTE["purple"], PALETTE["gray"]]
    edges = [PALETTE["blue_edge"], PALETTE["teal_edge"], PALETTE["green_edge"], PALETTE["amber_edge"], PALETTE["rose_edge"], PALETTE["purple_edge"], PALETTE["gray_edge"]]
    centers: list[tuple[int, int]] = []
    for i, node in enumerate(nodes):
        x = x0 + i * 360
        diagram.box((x, y, x + box_w, y + box_h), labels.get(node, node), fills[i], edges[i], font=FONT_BOOK.get(23))
        centers.append((x + box_w // 2, y + box_h // 2))
    for first, second in zip(centers, centers[1:]):
        diagram.arrow((first[0] + box_w // 2, first[1]), (second[0] - box_w // 2, second[1]))
    diagram.draw.rounded_rectangle((220, 505, 2380, 610), radius=18, fill="#F8FAFC", outline="#D6DDE6", width=2)
    diagram.draw.text(
        (1300, 558),
        "底座复用：模型接入、工具治理、知识记忆、协作调度、实时投影、多用户权限与多端发布",
        fill=PALETTE["ink"],
        font=FONT_BOOK.get(26),
        anchor="mm",
    )
    diagram.save(path)


def render_resource_model(figure: FigureBlock, path: Path) -> None:
    diagram = Diagram(2500, 1600, figure.title)
    node_labels = {
        "ORG_UNITS": "ORG_UNITS\n组织单元",
        "USER_ACCOUNTS": "USER_ACCOUNTS\n用户账户",
        "USER_AGENTS": "USER_AGENTS\n用户智能体",
        "CHAT_SESSIONS": "CHAT_SESSIONS\n会话",
        "HIVES": "HIVES\n协作空间",
        "MEMORY_RECORDS": "MEMORY_RECORDS\n记忆记录",
        "USER_TOKENS": "USER_TOKENS\n令牌",
        "USER_TOOL_ACCESS": "USER_TOOL_ACCESS\n工具授权",
        "AGENT_THREADS": "AGENT_THREADS\n智能体线程",
        "TEAM_RUNS": "TEAM_RUNS\n协作任务",
        "TEAM_TASKS": "TEAM_TASKS\n子任务",
        "STREAM_EVENTS": "STREAM_EVENTS\n事件流",
        "TOOL_LOGS": "TOOL_LOGS\n工具日志",
    }
    pos = {
        "ORG_UNITS": (930, 140, 1570, 245),
        "USER_ACCOUNTS": (930, 345, 1570, 470),
        "USER_AGENTS": (145, 660, 585, 775),
        "CHAT_SESSIONS": (665, 660, 1105, 775),
        "HIVES": (1185, 660, 1625, 775),
        "MEMORY_RECORDS": (1705, 660, 2145, 775),
        "USER_TOKENS": (665, 930, 1105, 1045),
        "USER_TOOL_ACCESS": (1185, 930, 1625, 1045),
        "AGENT_THREADS": (145, 1065, 585, 1180),
        "TEAM_RUNS": (1185, 1065, 1625, 1180),
        "TEAM_TASKS": (1185, 1300, 1625, 1415),
        "STREAM_EVENTS": (665, 1220, 1105, 1335),
        "TOOL_LOGS": (665, 1405, 1105, 1520),
    }
    for node, xy in pos.items():
        fill, edge = (PALETTE["blue"], PALETTE["blue_edge"]) if node in {"ORG_UNITS", "USER_ACCOUNTS"} else (PALETTE["gray"], PALETTE["gray_edge"])
        if node in {"HIVES", "TEAM_RUNS", "TEAM_TASKS"}:
            fill, edge = PALETTE["teal"], PALETTE["teal_edge"]
        if node in {"USER_TOOL_ACCESS", "TOOL_LOGS"}:
            fill, edge = PALETTE["rose"], PALETTE["rose_edge"]
        diagram.box(xy, node_labels[node], fill, edge, font=FONT_BOOK.get(22))
    def bottom(node: str) -> tuple[int, int]:
        x1, y1, x2, y2 = pos[node]
        return ((x1 + x2) // 2, y2)
    def top(node: str) -> tuple[int, int]:
        x1, y1, x2, y2 = pos[node]
        return ((x1 + x2) // 2, y1)
    def edge(a: str, b: str, label: str = "") -> None:
        diagram.arrow(bottom(a), top(b), label=label)
    edge("ORG_UNITS", "USER_ACCOUNTS", "contains")
    for node, label in [
        ("USER_AGENTS", "owns"),
        ("CHAT_SESSIONS", "owns"),
        ("HIVES", "owns"),
        ("MEMORY_RECORDS", "owns"),
        ("USER_TOKENS", "has"),
        ("USER_TOOL_ACCESS", "grants"),
    ]:
        sx = (pos["USER_ACCOUNTS"][0] + pos["USER_ACCOUNTS"][2]) // 2
        sy = pos["USER_ACCOUNTS"][3]
        tx = (pos[node][0] + pos[node][2]) // 2
        ty = pos[node][1]
        mid_y = 560 if pos[node][1] < 900 else 860
        diagram.elbow_arrow([(sx, sy), (sx, mid_y), (tx, mid_y), (tx, ty)], label=label)
    edge("USER_AGENTS", "AGENT_THREADS", "runs")
    edge("HIVES", "TEAM_RUNS", "contains")
    edge("TEAM_RUNS", "TEAM_TASKS", "contains")
    edge("CHAT_SESSIONS", "STREAM_EVENTS", "emits")
    edge("CHAT_SESSIONS", "TOOL_LOGS", "records")
    diagram.draw.text((1250, 1518), "资源归属先于调度，协作任务必须落在合法用户或组织边界内", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm")
    diagram.save(path)


def render_tool_governance(figure: FigureBlock, path: Path) -> None:
    labels = parse_flow_nodes(figure.code)
    diagram = Diagram(2600, 980, figure.title)
    row_y = 355
    nodes = ["Model", "Parse", "Visible", "Permission", "Approval", "Sandbox", "Execute", "Normalize", "Audit"]
    pos = {
        "Model": (80, row_y, 330, row_y + 120),
        "Parse": (410, row_y, 660, row_y + 120),
        "Visible": (740, row_y, 990, row_y + 120),
        "Permission": (1070, row_y, 1370, row_y + 120),
        "Approval": (1450, row_y - 25, 1710, row_y + 145),
        "Sandbox": (1790, row_y, 2040, row_y + 120),
        "Execute": (2120, row_y, 2370, row_y + 120),
        "Normalize": (1790, 660, 2040, 780),
        "Audit": (1450, 660, 1710, 780),
        "Wait": (1450, 150, 1710, 270),
    }
    for node, xy in pos.items():
        if node == "Approval":
            diagram.diamond(xy, labels.get(node, node))
        else:
            fill, edge = PALETTE["gray"], PALETTE["gray_edge"]
            if node in {"Visible", "Permission", "Wait"}:
                fill, edge = PALETTE["rose"], PALETTE["rose_edge"]
            if node in {"Sandbox", "Execute"}:
                fill, edge = PALETTE["amber"], PALETTE["amber_edge"]
            if node in {"Normalize", "Audit"}:
                fill, edge = PALETTE["green"], PALETTE["green_edge"]
            diagram.box(xy, labels.get(node, node), fill, edge, font=FONT_BOOK.get(23))
    for a, b in zip(nodes[:4], nodes[1:5]):
        diagram.arrow(((pos[a][2]), (pos[a][1] + pos[a][3]) // 2), ((pos[b][0]), (pos[b][1] + pos[b][3]) // 2))
    diagram.arrow((1710, 415), (1790, 415), "否")
    diagram.elbow_arrow([(1580, 330), (1580, 270)], "是")
    diagram.elbow_arrow([(1580, 270), (1580, 330), (1835, 330), (1835, 355)])
    diagram.arrow((2040, 415), (2120, 415))
    diagram.elbow_arrow([(2245, 475), (2245, 720), (2040, 720)])
    diagram.arrow((1790, 720), (1710, 720))
    diagram.elbow_arrow([(1450, 720), (205, 720), (205, 475)], "记录后回到模型上下文")
    diagram.draw.text((1300, 875), "可见性、授权、审批、隔离、审计均在运行时完成，不依赖模型自觉遵守", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm")
    diagram.save(path)


def render_sequence(figure: FigureBlock, path: Path) -> None:
    participants, messages = parse_sequence(figure.code)
    height = 1420 if figure.index == 6 else 1240
    diagram = Diagram(2400, height, figure.title)
    x0 = 180
    x_gap = (diagram.width - 2 * x0) // (len(participants) - 1)
    xs = {pid: x0 + i * x_gap for i, (pid, _) in enumerate(participants)}
    top = 150
    bottom = height - 135
    for pid, label in participants:
        x = xs[pid]
        diagram.box((x - 125, top, x + 125, top + 80), label, PALETTE["blue"], PALETTE["blue_edge"], font=FONT_BOOK.get(24))
        diagram.dashed_line((x, top + 76), (x, bottom), fill="#B6C0CC", width=2, dash=14, gap=12)
    if figure.index == 6:
        y_positions = [260, 390, 470, 590, 690, 805, 930, 1045, 1170, 1290]
    else:
        y_positions = [265, 410, 515, 670, 775, 900, 1015, 1135]
    labels_above = True
    for idx, (src, dst, label) in enumerate(messages):
        if src not in xs or dst not in xs:
            continue
        sx, dx = xs[src], xs[dst]
        y_pos = y_positions[min(idx, len(y_positions) - 1)]
        if src == dst:
            direction = 1 if sx < diagram.width - 360 else -1
            elbow = sx + direction * 145
            diagram.elbow_arrow([(sx, y_pos), (elbow, y_pos), (elbow, y_pos + 48), (sx, y_pos + 48)])
            diagram.label((elbow, y_pos - 28), label)
        else:
            start = (sx + (34 if dx > sx else -34), y_pos)
            end = (dx - (34 if dx > sx else -34), y_pos)
            diagram.arrow(start, end)
            label_y = y_pos - 42 if labels_above else y_pos + 42
            diagram.label(((sx + dx) // 2, label_y), label)
            labels_above = not labels_above
    if figure.index == 6:
        diagram.draw.rounded_rectangle((620, 540, 2020, 1110), radius=18, outline=PALETTE["amber_edge"], width=3)
        diagram.draw.text((650, 560), "模型轮次循环：模型生成、工具治理、事件写入、最终回复", fill=PALETTE["amber_edge"], font=FONT_BOOK.get(24), anchor="la")
    diagram.save(path)


def render_turn_model(figure: FigureBlock, path: Path) -> None:
    diagram = Diagram(2400, 1320, figure.title)
    top_boxes = [
        ((120, 170, 480, 290), "用户提交\n一次输入", PALETTE["blue"], PALETTE["blue_edge"]),
        ((610, 170, 970, 290), "线程运行时\n获取执行租约", PALETTE["blue"], PALETTE["blue_edge"]),
        ((1100, 170, 1460, 290), "执行内核\n启动用户轮次", PALETTE["amber"], PALETTE["amber_edge"]),
        ((1590, 170, 1950, 290), "事件流\n记录过程", PALETTE["purple"], PALETTE["purple_edge"]),
    ]
    for xy, text, fill, edge in top_boxes:
        diagram.box(xy, text, fill, edge, font=FONT_BOOK.get(27))
    for first, second in zip(top_boxes, top_boxes[1:3]):
        diagram.arrow((first[0][2], (first[0][1] + first[0][3]) // 2), (second[0][0], (second[0][1] + second[0][3]) // 2))
    diagram.elbow_arrow([(1280, 290), (1280, 380)])

    loop = (230, 380, 2170, 1010)
    diagram.draw.rounded_rectangle(loop, radius=28, fill="#FFFDF7", outline=PALETTE["amber_edge"], width=4)
    diagram.draw.text((270, 420), "模型轮次循环", fill=PALETTE["amber_edge"], font=FONT_BOOK.get(30, bold=True), anchor="la")
    diagram.draw.text((270, 462), "一个用户轮次可以包含多次模型调用、工具执行、事件写入和重试。", fill=PALETTE["muted"], font=FONT_BOOK.get(24), anchor="la")

    boxes = {
        "context": ((340, 555, 730, 690), "组装上下文\n冻结提示词 + 记忆快照", PALETTE["blue"], PALETTE["blue_edge"]),
        "model": ((880, 555, 1270, 690), "请求模型\n生成文本或工具调用", PALETTE["amber"], PALETTE["amber_edge"]),
        "tool": ((1420, 555, 1810, 690), "工具治理\n权限、审批、沙盒", PALETTE["rose"], PALETTE["rose_edge"]),
        "result": ((880, 800, 1270, 935), "工具结果\n进入下一模型轮次", PALETTE["teal"], PALETTE["teal_edge"]),
        "final": ((1420, 800, 1810, 935), "最终回复\n写入最终事件", PALETTE["green"], PALETTE["green_edge"]),
    }
    for xy, text, fill, edge in boxes.values():
        diagram.box(xy, text, fill, edge, font=FONT_BOOK.get(25))
    diagram.arrow((730, 622), (880, 622))
    diagram.arrow((1270, 622), (1420, 622), "需要工具")
    diagram.arrow((1810, 622), (1960, 622))
    diagram.elbow_arrow([(1960, 622), (1960, 868), (1810, 868)], "工具返回")
    diagram.arrow((1270, 868), (1420, 868), "生成最终回复")
    diagram.elbow_arrow([(880, 868), (520, 868), (520, 690)], "继续轮次", dashed=True)
    diagram.elbow_arrow([(1615, 800), (1615, 735), (1770, 735), (1770, 690)], "阶段事件")
    diagram.elbow_arrow([(1615, 935), (1615, 1070)])

    bottom_boxes = [
        ((1210, 1070, 1570, 1190), "用户轮次完成", PALETTE["gray"], PALETTE["gray_edge"]),
        ((1700, 1070, 2060, 1190), "返回结果\n或继续投影", PALETTE["purple"], PALETTE["purple_edge"]),
    ]
    for xy, text, fill, edge in bottom_boxes:
        diagram.box(xy, text, fill, edge, font=FONT_BOOK.get(26))
    diagram.arrow((1570, 1130), (1700, 1130))
    diagram.elbow_arrow([(1950, 230), (2070, 230), (2070, 1130), (2060, 1130)], "过程可观察")
    diagram.draw.text((1200, 1240), "用户轮次描述一次完整请求；模型轮次描述请求内部的每次模型调用和工具循环。", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm")
    diagram.save(path)


def render_event_replay(figure: FigureBlock, path: Path) -> None:
    diagram = Diagram(2400, 1240, figure.title)
    diagram.draw.rounded_rectangle((100, 150, 2300, 510), radius=26, fill="#FBFCFE", outline="#D6DDE6", width=3)
    diagram.draw.text((145, 172), "运行事件发布链路", fill=PALETTE["muted"], font=FONT_BOOK.get(26, bold=True), anchor="la")
    publish = [
        ((160, 270, 465, 390), "运行时\n发布运行事件", PALETTE["amber"], PALETTE["amber_edge"]),
        ((575, 270, 880, 390), "事件服务\n生成递增编号", PALETTE["blue"], PALETTE["blue_edge"]),
        ((990, 270, 1295, 390), "存储\n持久化事件", PALETTE["green"], PALETTE["green_edge"]),
        ((1405, 270, 1710, 390), "实时连接\n广播事件", PALETTE["purple"], PALETTE["purple_edge"]),
        ((1820, 270, 2125, 390), "客户端\n接收并应用", PALETTE["teal"], PALETTE["teal_edge"]),
    ]
    for xy, text, fill, edge in publish:
        diagram.box(xy, text, fill, edge, font=FONT_BOOK.get(25))
    for first, second in zip(publish, publish[1:]):
        diagram.arrow((first[0][2], 330), (second[0][0], 330))

    diagram.draw.rounded_rectangle((100, 610, 2300, 1080), radius=26, fill="#FBFCFE", outline="#D6DDE6", width=3)
    diagram.draw.text((145, 632), "断线补水与重放链路", fill=PALETTE["muted"], font=FONT_BOOK.get(26, bold=True), anchor="la")
    replay = [
        ((160, 760, 465, 900), "客户端重连\n携带最后事件编号", PALETTE["teal"], PALETTE["teal_edge"]),
        ((575, 760, 880, 900), "实时连接\n请求缺失事件", PALETTE["purple"], PALETTE["purple_edge"]),
        ((990, 760, 1295, 900), "事件服务\n计算事件区间", PALETTE["blue"], PALETTE["blue_edge"]),
        ((1405, 760, 1710, 900), "存储\n查询事件列表", PALETTE["green"], PALETTE["green_edge"]),
        ((1820, 760, 2125, 900), "客户端\n去重、应用、收敛", PALETTE["teal"], PALETTE["teal_edge"]),
    ]
    for xy, text, fill, edge in replay:
        diagram.box(xy, text, fill, edge, font=FONT_BOOK.get(25))
    for first, second in zip(replay, replay[1:]):
        diagram.arrow((first[0][2], 830), (second[0][0], 830))
    diagram.elbow_arrow([(1558, 760), (1558, 690), (1150, 690), (1150, 760)], "返回缺失事件", color=PALETTE["line"])
    diagram.elbow_arrow([(1558, 900), (1558, 1000), (1972, 1000), (1972, 900)], "重放事件", color=PALETTE["line"])
    diagram.draw.text((1200, 1140), "事件编号是投影恢复的锚点；前端只根据事件和快照收敛，不反向改写运行时真相。", fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm")
    diagram.save(path)


def render_collaboration(figure: FigureBlock, path: Path) -> None:
    labels = parse_flow_nodes(figure.code)
    diagram = Diagram(2500, 1500, figure.title)
    pos = {
        "Goal": (165, 180, 585, 285),
        "Boundary": (165, 375, 585, 495),
        "Plan": (165, 585, 585, 705),
        "Select": (165, 795, 585, 915),
        "Guard": (110, 1005, 640, 1135),
        "Dispatch": (165, 1230, 585, 1340),
        "A": (870, 395, 1290, 505),
        "B": (870, 665, 1290, 775),
        "C": (870, 935, 1290, 1045),
        "Report": (1545, 665, 2015, 785),
        "Merge": (1545, 905, 2015, 1025),
        "Board": (1545, 1145, 2015, 1265),
        "User": (1545, 1325, 2015, 1435),
    }
    diagram.draw.rounded_rectangle((80, 145, 675, 1375), radius=26, fill="#FBFCFE", outline="#D6DDE6", width=3)
    diagram.draw.rounded_rectangle((775, 290, 1385, 1120), radius=26, fill="#FBFCFE", outline="#D6DDE6", width=3)
    diagram.draw.rounded_rectangle((1475, 560, 2090, 1460), radius=26, fill="#FBFCFE", outline="#D6DDE6", width=3)
    diagram.draw.text((375, 150), "调度控制链路", fill=PALETTE["muted"], font=FONT_BOOK.get(24), anchor="mt")
    diagram.draw.text((1080, 305), "并行执行区", fill=PALETTE["muted"], font=FONT_BOOK.get(24), anchor="mt")
    diagram.draw.text((1780, 575), "汇总与投影区", fill=PALETTE["muted"], font=FONT_BOOK.get(24), anchor="mt")
    for node, xy in pos.items():
        fill, edge = PALETTE["gray"], PALETTE["gray_edge"]
        if node in {"Goal", "User"}:
            fill, edge = PALETTE["blue"], PALETTE["blue_edge"]
        if node in {"Boundary", "Guard"}:
            fill, edge = PALETTE["rose"], PALETTE["rose_edge"]
        if node in {"A", "B", "C", "Dispatch", "Report"}:
            fill, edge = PALETTE["teal"], PALETTE["teal_edge"]
        if node in {"Merge", "Board"}:
            fill, edge = PALETTE["green"], PALETTE["green_edge"]
        diagram.box(xy, labels.get(node, node), fill, edge, font=FONT_BOOK.get(23))
    def center(node: str) -> tuple[int, int]:
        x1, y1, x2, y2 = pos[node]
        return ((x1 + x2) // 2, (y1 + y2) // 2)
    for a, b in [("Goal", "Boundary"), ("Boundary", "Plan"), ("Plan", "Select"), ("Select", "Guard"), ("Guard", "Dispatch")]:
        diagram.arrow((center(a)[0], pos[a][3]), (center(b)[0], pos[b][1]))
    bus_x = 735
    diagram.arrow((pos["Dispatch"][2], center("Dispatch")[1]), (bus_x, center("Dispatch")[1]))
    for node in ["A", "B", "C"]:
        y = center(node)[1]
        diagram.elbow_arrow([(bus_x, center("Dispatch")[1]), (bus_x, y), (pos[node][0], y)])
        diagram.arrow((pos[node][2], y), (pos["Report"][0], y))
    diagram.elbow_arrow([(center("Report")[0], pos["Report"][3]), (center("Merge")[0], pos["Merge"][1])])
    diagram.elbow_arrow([(center("Merge")[0], pos["Merge"][3]), (center("Board")[0], pos["Board"][1])])
    diagram.elbow_arrow([(center("Board")[0], pos["Board"][3]), (center("User")[0], pos["User"][1])])
    diagram.draw.rounded_rectangle((775, 1220, 1385, 1370), radius=18, fill="#F8FAFC", outline="#D6DDE6", width=2)
    note = diagram.wrap_text("策略守卫控制同空间、并发上限、递归深度、超时与重试，避免越界和失控调度", 540, diagram.caption_font)
    diagram.draw.multiline_text((1080, 1295), note, fill=PALETTE["muted"], font=diagram.caption_font, anchor="mm", align="center", spacing=8)
    diagram.save(path)


def render_figure(figure: FigureBlock, path: Path) -> None:
    if figure.index == 1:
        render_state_model(figure, path)
    elif figure.index == 2:
        render_architecture(figure, path)
    elif figure.index == 3:
        render_linear_path(figure, path)
    elif figure.index == 4:
        render_resource_model(figure, path)
    elif figure.index == 5:
        render_tool_governance(figure, path)
    elif figure.index == 6:
        render_turn_model(figure, path)
    elif figure.index == 8:
        render_event_replay(figure, path)
    elif figure.index == 7:
        render_collaboration(figure, path)
    else:
        fallback = Diagram(1600, 900, figure.title)
        fallback.draw.text((800, 450), clean_label(figure.code[:500]), fill=PALETTE["ink"], font=fallback.node_font, anchor="mm")
        fallback.save(path)


def disable_snap_to_grid(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "0")


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def normalize_picture_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(10)
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_snap_to_grid(paragraph)


def normalize_caption_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(4)
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_snap_to_grid(paragraph)
    for run in paragraph.runs:
        run.font.bold = True


def replace_image_placeholders(docx_path: Path, assets_dir: Path) -> int:
    doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)
    replaced = 0
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text
        match = re.search(r"assets[/\\]fig(\d+)\.png", text)
        if not match:
            continue
        image_path = assets_dir / f"fig{match.group(1)}.png"
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        clear_paragraph(paragraph)
        normalize_picture_paragraph(paragraph)
        paragraph.add_run().add_picture(str(image_path), width=Cm(15.2))
        replaced += 1
        if index > 0 and re.match(r"^\s*图\s*[0-9０-９]+", paragraphs[index - 1].text.strip()):
            normalize_caption_paragraph(paragraphs[index - 1])
    doc.save(docx_path)
    return replaced


def run_converter() -> None:
    converter = REPO_ROOT / "config" / "skills" / "公文写作" / "scripts" / "convert_markdown_to_docx.py"
    if not converter.exists():
        raise FileNotFoundError(converter)
    command = [
        sys.executable,
        str(converter),
        str(SOURCE_MD_PATH),
        "--output",
        str(DOCX_PATH),
        "--image-width",
        "15.2cm",
    ]
    subprocess.run(command, cwd=REPO_ROOT, check=True)


def validate_docx(docx_path: Path) -> dict[str, int | bool]:
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    with zipfile.ZipFile(docx_path) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "embedded_media": len(media),
        "contains_mermaid": "mermaid" in text or "```" in text,
        "contains_placeholder": "图片：" in text,
        "contains_mojibake": "???" in text or "???" in document_xml,
    }


def non_white_ratio(path: Path) -> float:
    image = Image.open(path).convert("RGB")
    resized = image.resize((240, max(1, int(240 * image.height / image.width))))
    pixel_bytes = resized.tobytes()
    total = 0
    non_white = 0
    for offset in range(0, len(pixel_bytes), 3):
        r, g, b = pixel_bytes[offset], pixel_bytes[offset + 1], pixel_bytes[offset + 2]
        total += 1
        if min(255 - r, 255 - g, 255 - b) < 0:
            continue
        if (r, g, b) < (248, 248, 248):
            non_white += 1
    return non_white / max(1, total)


def build_contact_sheet(paths: Iterable[Path]) -> None:
    thumbs: list[Image.Image] = []
    for path in paths:
        img = Image.open(path).convert("RGB")
        img.thumbnail((600, 360), Image.Resampling.LANCZOS)
        canvas = Image.new("RGB", (640, 430), "white")
        draw = ImageDraw.Draw(canvas)
        canvas.paste(img, ((640 - img.width) // 2, 46))
        draw.text((320, 20), path.stem, fill=PALETTE["ink"], font=FONT_BOOK.get(24), anchor="mm")
        thumbs.append(canvas)
    cols = 2
    rows = math.ceil(len(thumbs) / cols)
    sheet = Image.new("RGB", (cols * 640, rows * 430), "#F4F6F8")
    for idx, thumb in enumerate(thumbs):
        sheet.paste(thumb, ((idx % cols) * 640, (idx // cols) * 430))
    CONTACT_SHEET_PATH.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(CONTACT_SHEET_PATH, "PNG", dpi=(180, 180), optimize=True)


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    paper_path = locate_paper_markdown()
    markdown = paper_path.read_text(encoding="utf-8-sig")
    figures = extract_figures(markdown)
    if len(figures) != 8:
        raise RuntimeError(f"Expected 8 Mermaid figures, found {len(figures)}.")

    image_paths: list[Path] = []
    for figure in figures:
        image_path = ASSETS_DIR / f"fig{figure.index}.png"
        render_figure(figure, image_path)
        ratio = non_white_ratio(image_path)
        if ratio < 0.01:
            raise RuntimeError(f"Generated figure appears blank: {image_path}")
        image_paths.append(image_path)

    build_contact_sheet(image_paths)
    SOURCE_MD_PATH.write_text(build_source_markdown(markdown, figures), encoding="utf-8")
    run_converter()
    replaced = replace_image_placeholders(DOCX_PATH, ASSETS_DIR)
    if replaced != len(figures):
        raise RuntimeError(f"Expected to replace {len(figures)} image placeholders, replaced {replaced}.")

    report = validate_docx(DOCX_PATH)
    if report["embedded_media"] != len(figures):
        raise RuntimeError(f"Expected {len(figures)} embedded media files, found {report['embedded_media']}.")
    if report["contains_mermaid"] or report["contains_placeholder"] or report["contains_mojibake"]:
        raise RuntimeError(f"DOCX validation failed: {report}")

    print(f"source={paper_path}")
    print(f"docx={DOCX_PATH}")
    print(f"assets={ASSETS_DIR}")
    print(f"contact_sheet={CONTACT_SHEET_PATH}")
    print(f"validation={report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
